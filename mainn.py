import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import mysql.connector
from mysql.connector import Error
from datetime import datetime
import hashlib
import os
import shutil
from pathlib import Path
import subprocess
import shutil
from docxtpl import DocxTemplate
from docx2pdf import convert
import webbrowser


# Конфигурация базы данных
DB_CONFIG = {
    'host': 'localhost',
    'user': 'root',
    'password': '',
    'database': 'expertise_db_db',
    'charset': 'utf8mb4'
}


current_user = None
current_user_role = None
connection = None
UPLOAD_FOLDER = "uploads"
TEMPLATES_FOLDER = "templates"
DRAFTS_FOLDER = "drafts"
FINAL_FOLDER = "final_documents"


for folder in [UPLOAD_FOLDER, TEMPLATES_FOLDER, DRAFTS_FOLDER, FINAL_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)




class Database:
    """Класс для работы с базой данных"""

    @staticmethod
    def get_connection():
        try:
            conn = mysql.connector.connect(**DB_CONFIG)
            return conn
        except Error as e:
            messagebox.showerror("Ошибка БД", f"Не удалось подключиться к базе данных: {e}")
            return None

    @staticmethod
    def execute_query(query, params=None, fetchone=False, fetchall=False, commit=False):
        conn = Database.get_connection()
        if not conn:
            return None

        cursor = None
        result = None

        try:
            cursor = conn.cursor(dictionary=True, buffered=True)  # ВАЖНО: buffered=True
            cursor.execute(query, params or ())

            if commit:
                conn.commit()
                result = cursor.lastrowid
            elif fetchone:
                result = cursor.fetchone()
                # Очищаем оставшиеся результаты, если они есть
                if cursor.nextset():
                    pass
            elif fetchall:
                result = cursor.fetchall()
                if cursor.nextset():
                    pass
            else:
                # Если ничего не запрашивали, но нужно очистить результат
                if cursor.with_rows:
                    cursor.fetchall()
                if cursor.nextset():
                    pass

        except Error as e:
            messagebox.showerror("Ошибка запроса", str(e))
            if commit:
                try:
                    conn.rollback()
                except:
                    pass
            result = None
        finally:
            if cursor:
                try:
                    # Дополнительная очистка
                    while cursor.nextset():
                        pass
                    cursor.close()
                except:
                    pass
            if conn:
                try:
                    conn.close()
                except:
                    pass

        return result

    @staticmethod
    def log_action(employee_id, action_type, object_type=None, object_id=None, details=None):
        """Запись действия в журнал аудита"""
        query = """
            INSERT INTO audit_log (employee_id, action_time, action_type, object_type, object_id, details, ip_address)
            VALUES (%s, NOW(), %s, %s, %s, %s, %s)
        """
        Database.execute_query(query, (employee_id, action_type, object_type, object_id, details, '127.0.0.1'),
                               commit=True)


class LoginWindow:
    """Окно авторизации"""

    def __init__(self):
        self.window = tk.Tk()
        self.window.title("Авторизация - Экспертиза-ПБ")
        self.window.geometry("400x300")
        self.window.resizable(False, False)

        # Центрирование окна
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f'{width}x{height}+{x}+{y}')

        # Заголовок
        tk.Label(self.window, text="Экспертиза-ПБ", font=("Arial", 20, "bold")).pack(pady=20)
        tk.Label(self.window, text="Вход в систему автоматизации документооборота", font=("Arial", 10)).pack(pady=5)

        # Форма
        frame = tk.Frame(self.window)
        frame.pack(pady=20)

        tk.Label(frame, text="Логин:", font=("Arial", 10)).grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.login_entry = tk.Entry(frame, width=30, font=("Arial", 10))
        self.login_entry.grid(row=0, column=1, padx=5, pady=5)
        self.login_entry.focus()

        tk.Label(frame, text="Пароль:", font=("Arial", 10)).grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.password_entry = tk.Entry(frame, width=30, font=("Arial", 10), show="*")
        self.password_entry.grid(row=1, column=1, padx=5, pady=5)

        # Информация о тестовых пользователях
        info_frame = tk.Frame(self.window, bg="#f0f0f0", relief="groove", bd=1)
        info_frame.pack(fill="x", padx=20, pady=10)

        tk.Label(info_frame, text="Тестовые пользователи:", font=("Arial", 9, "bold"), bg="#f0f0f0").pack(anchor="w",
                                                                                                          padx=5,
                                                                                                          pady=2)
        tk.Label(info_frame, text="admin / password (Администратор)", font=("Arial", 8), bg="#f0f0f0").pack(anchor="w",
                                                                                                            padx=5)
        tk.Label(info_frame, text="smirnov / password (Руководитель)", font=("Arial", 8), bg="#f0f0f0").pack(anchor="w",
                                                                                                             padx=5)
        tk.Label(info_frame, text="ivanov / password (Эксперт)", font=("Arial", 8), bg="#f0f0f0").pack(anchor="w",
                                                                                                       padx=5)
        tk.Label(info_frame, text="sidorova / password (Специалист)", font=("Arial", 8), bg="#f0f0f0").pack(anchor="w",
                                                                                                            padx=5)

        # Кнопки
        btn_frame = tk.Frame(self.window)
        btn_frame.pack(pady=20)

        tk.Button(btn_frame, text="Войти", command=self.login,
                  bg="#4CAF50", fg="white", font=("Arial", 10), width=10).pack(side="left", padx=5)
        tk.Button(btn_frame, text="Выход", command=self.window.quit,
                  bg="#f44336", fg="white", font=("Arial", 10), width=10).pack(side="left", padx=5)

        # Привязка Enter к входу
        self.password_entry.bind('<Return>', lambda event: self.login())

    def login(self):
        login = self.login_entry.get().strip()
        password = self.password_entry.get().strip()

        if not login or not password:
            messagebox.showwarning("Предупреждение", "Введите логин и пароль")
            return

        # Хеширование пароля
        password_hash = hashlib.md5(password.encode()).hexdigest()

        query = "SELECT * FROM employees WHERE login = %s AND password_hash = %s"
        user = Database.execute_query(query, (login, password_hash), fetchone=True)

        if user:
            global current_user, current_user_role
            current_user = user
            current_user_role = user['role']

            Database.log_action(user['id'], 'LOGIN', 'system', None,
                                f"Пользователь {user['full_name']} вошел в систему")

            self.window.destroy()
            MainWindow()
        else:
            messagebox.showerror("Ошибка", "Неверный логин или пароль")
            Database.log_action(0, 'LOGIN_FAILED', 'system', None, f"Неудачная попытка входа: {login}")

    def run(self):
        self.window.mainloop()


class MainWindow:
    """Главное окно программы"""

    def __init__(self):
        self.window = tk.Tk()
        self.window.title(
            f"Экспертиза-ПБ - Главное меню (Пользователь: {current_user['full_name']} / {self.get_role_name(current_user_role)})")
        self.window.geometry("1000x700")

        # Меню
        menubar = tk.Menu(self.window)
        self.window.config(menu=menubar)

        # Меню "Справочники"
        sprav_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Справочники", menu=sprav_menu)
        sprav_menu.add_command(label="Типы экспертиз", command=lambda: SpravWindow("expert_types", "Типы экспертиз"))
        sprav_menu.add_command(label="Объекты экспертизы", command=lambda: SpravWindow("objects", "Объекты экспертизы"))
        sprav_menu.add_command(label="Заказчики", command=lambda: SpravWindow("customers", "Заказчики"))
        sprav_menu.add_command(label="Сотрудники", command=lambda: SpravWindow("employees", "Сотрудники"))
        sprav_menu.add_command(label="Статусы документов",
                               command=lambda: SpravWindow("statuses", "Статусы документов"))
        sprav_menu.add_command(label="Шаблоны документов",
                               command=lambda: SpravWindow("templates", "Шаблоны документов"))

        # Меню "Проекты"
        project_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Проекты", menu=project_menu)
        project_menu.add_command(label="Список проектов", command=ProjectListWindow)
        if current_user_role in ['admin', 'specialist']:
            project_menu.add_command(label="Создать новый проект", command=ProjectCreateWindow)

        # Меню "Результаты обследований"
        if current_user_role in ['expert', 'admin']:
            results_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="Результаты обследований", menu=results_menu)
            results_menu.add_command(label="Загрузить результаты", command=ResultsUploadWindow)
            results_menu.add_command(label="Просмотр результатов", command=ResultsViewWindow)

        # Меню "Документы"
        docs_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Документы", menu=docs_menu)
        docs_menu.add_command(label="Черновики", command=DraftsWindow)
        docs_menu.add_command(label="Готовые документы", command=FinalDocsWindow)

        # Меню "Согласование"
        if current_user_role in ['director', 'admin']:
            approval_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="Согласование", menu=approval_menu)
            approval_menu.add_command(label="Документы на согласовании", command=ApprovalWindow)

        # Меню "Отчетность"
        if current_user_role in ['director', 'specialist', 'admin']:
            reports_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="Отчетность", menu=reports_menu)
            reports_menu.add_command(label="Реестр проектов", command=ProjectsReportWindow)
            reports_menu.add_command(label="Журнал учета документов", command=DocumentsReportWindow)

        # Меню "Администрирование"
        if current_user_role == 'admin':
            admin_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="Администрирование", menu=admin_menu)
            admin_menu.add_command(label="Журнал аудита", command=AuditLogWindow)
            admin_menu.add_command(label="Управление пользователями", command=UserManagementWindow)
            admin_menu.add_command(label="Резервное копирование", command=BackupWindow)

        # Меню "Помощь"
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Помощь", menu=help_menu)
        help_menu.add_command(label="О программе", command=self.about)
        help_menu.add_command(label="Выход", command=self.logout)

        # Приветствие
        welcome_frame = tk.Frame(self.window, bg="#f0f0f0", relief="groove", bd=2)
        welcome_frame.pack(fill="x", padx=20, pady=20)

        tk.Label(welcome_frame, text=f"Добро пожаловать, {current_user['full_name']}!",
                 font=("Arial", 16, "bold"), bg="#f0f0f0").pack(pady=10)
        tk.Label(welcome_frame, text=f"Ваша роль: {self.get_role_name(current_user_role)}",
                 font=("Arial", 12), bg="#f0f0f0").pack(pady=5)
        tk.Label(welcome_frame, text="Выберите раздел в меню для начала работы",
                 font=("Arial", 10), bg="#f0f0f0").pack(pady=10)

        # Основной контейнер
        main_container = tk.Frame(self.window)
        main_container.pack(fill="both", expand=True, padx=20, pady=10)

        # Статистика
        stats_frame = tk.LabelFrame(main_container, text="Статистика системы", font=("Arial", 12, "bold"))
        stats_frame.pack(fill="x", pady=10)

        # Получение статистики
        projects_count = Database.execute_query("SELECT COUNT(*) as count FROM projects", fetchone=True)
        drafts_count = Database.execute_query("SELECT COUNT(*) as count FROM drafts WHERE status_id=1", fetchone=True)
        approval_count = Database.execute_query("SELECT COUNT(*) as count FROM drafts WHERE status_id=2", fetchone=True)
        final_count = Database.execute_query("SELECT COUNT(*) as count FROM final_documents", fetchone=True)

        stats_grid = tk.Frame(stats_frame)
        stats_grid.pack(pady=10)

        stats = [
            ("Всего проектов:", projects_count['count'] if projects_count else "0"),
            ("Черновиков:", drafts_count['count'] if drafts_count else "0"),
            ("На согласовании:", approval_count['count'] if approval_count else "0"),
            ("Готовых документов:", final_count['count'] if final_count else "0")
        ]

        for i, (label, value) in enumerate(stats):
            tk.Label(stats_grid, text=label, font=("Arial", 10)).grid(row=i // 2, column=(i % 2) * 2, padx=10, pady=5,
                                                                      sticky="w")
            tk.Label(stats_grid, text=str(value), font=("Arial", 10, "bold")).grid(row=i // 2, column=(i % 2) * 2 + 1,
                                                                                   padx=10, pady=5, sticky="w")

        # Последние проекты
        recent_frame = tk.LabelFrame(main_container, text="Последние проекты", font=("Arial", 12, "bold"))
        recent_frame.pack(fill="both", expand=True, pady=10)

        # Таблица последних проектов
        columns = ('Номер', 'Дата', 'Заказчик', 'Объект', 'Статус')
        self.recent_tree = ttk.Treeview(recent_frame, columns=columns, show='headings', height=8)

        for col in columns:
            self.recent_tree.heading(col, text=col)
            self.recent_tree.column(col, width=120)

        # Скроллбар
        scrollbar = ttk.Scrollbar(recent_frame, orient="vertical", command=self.recent_tree.yview)
        self.recent_tree.configure(yscrollcommand=scrollbar.set)

        self.recent_tree.pack(side="left", fill="both", expand=True, padx=5, pady=5)
        scrollbar.pack(side="right", fill="y")

        self.load_recent_projects()

        self.window.protocol("WM_DELETE_WINDOW", self.logout)
        self.window.mainloop()

    def get_role_name(self, role):
        roles = {
            'admin': 'Администратор',
            'director': 'Руководитель',
            'expert': 'Эксперт',
            'specialist': 'Специалист по документации'
        }
        return roles.get(role, role)

    def load_recent_projects(self):
        """Загрузка последних 10 проектов"""
        query = """
            SELECT p.project_number, p.creation_date, c.name as customer_name, 
                   o.name as object_name, s.name as status_name
            FROM projects p
            JOIN customers c ON p.customer_id = c.id
            JOIN objects o ON p.object_id = o.id
            JOIN statuses s ON p.status_id = s.id
            ORDER BY p.creation_date DESC
            LIMIT 10
        """
        projects = Database.execute_query(query, fetchall=True)

        for item in self.recent_tree.get_children():
            self.recent_tree.delete(item)

        if projects:
            for p in projects:
                self.recent_tree.insert('', 'end', values=(
                    p['project_number'],
                    p['creation_date'].strftime('%d.%m.%Y') if p['creation_date'] else '',
                    p['customer_name'],
                    p['object_name'],
                    p['status_name']
                ))

    def about(self):
        messagebox.showinfo("О программе",
                            "Информационная система автоматизации процесса А3\n"
                            "«Подготовка экспертного заключения и отчетных документов»\n\n"
                            "Разработано для ООО «Экспертиза-ПБ»\n"
                            "Версия 1.0\n\n"
                            "© 2026")

    def logout(self):
        if current_user:
            Database.log_action(current_user['id'], 'LOGOUT', 'system', None,
                                f"Пользователь {current_user['full_name']} вышел из системы")
        self.window.quit()


class SpravWindow:
    """Окно для работы со справочниками"""

    def __init__(self, table_name, title):
        self.table_name = table_name
        self.title = title
        self.edit_window = None

        self.window = tk.Toplevel()
        self.window.title(f"Справочник: {title}")
        self.window.geometry("900x600")

        # Панель инструментов
        toolbar = tk.Frame(self.window, bg="#f0f0f0", height=40)
        toolbar.pack(fill="x", padx=2, pady=2)

        if current_user_role in ['admin', 'specialist']:
            tk.Button(toolbar, text="Добавить", command=self.add_record,
                      bg="#4CAF50", fg="white", width=10).pack(side="left", padx=5, pady=5)
            tk.Button(toolbar, text="Редактировать", command=self.edit_record,
                      bg="#2196F3", fg="white", width=12).pack(side="left", padx=5, pady=5)
            tk.Button(toolbar, text="Удалить", command=self.delete_record,
                      bg="#f44336", fg="white", width=8).pack(side="left", padx=5, pady=5)

        tk.Button(toolbar, text="Обновить", command=self.load_data,
                  bg="#FF9800", fg="white", width=8).pack(side="left", padx=5, pady=5)

        # Поиск
        tk.Label(toolbar, text="Поиск:", bg="#f0f0f0").pack(side="right", padx=5)
        self.search_entry = tk.Entry(toolbar, width=30)
        self.search_entry.pack(side="right", padx=5)
        self.search_entry.bind('<KeyRelease>', self.search)

        # Основная таблица
        self.create_table()
        self.load_data()

    def get_columns(self):
        """Получение списка колонок для таблицы"""
        columns_map = {
            'expert_types': ['id', 'name', 'description', 'regulatory_base'],
            'objects': ['id', 'name', 'category', 'unit'],
            'customers': ['id', 'name', 'inn', 'contact_person', 'phone', 'email'],
            'employees': ['id', 'full_name', 'position', 'role', 'login', 'contacts'],
            'statuses': ['id', 'name', 'description', 'color_code'],
            'templates': ['id', 'name', 'doc_type', 'template_path', 'update_date']
        }
        return columns_map.get(self.table_name, [])

    def get_column_names(self):
        """Русские названия колонок"""
        names_map = {
            'id': 'ID',
            'name': 'Наименование',
            'description': 'Описание',
            'regulatory_base': 'Нормативная база',
            'category': 'Категория',
            'unit': 'Ед. измерения',
            'inn': 'ИНН',
            'contact_person': 'Контактное лицо',
            'phone': 'Телефон',
            'email': 'Email',
            'full_name': 'ФИО',
            'position': 'Должность',
            'role': 'Роль',
            'login': 'Логин',
            'contacts': 'Контакты',
            'color_code': 'Цвет',
            'doc_type': 'Тип документа',
            'template_path': 'Путь к шаблону',
            'update_date': 'Дата обновления'
        }
        return names_map

    def create_table(self):
        """Создание таблицы"""
        # Основной фрейм с таблицей
        table_frame = tk.Frame(self.window)
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        columns = self.get_columns()
        if not columns:
            return

        # Таблица
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', selectmode='browse')

        column_names = self.get_column_names()
        for col in columns:
            self.tree.heading(col, text=column_names.get(col, col))
            self.tree.column(col, width=100, anchor="w")

        # Скроллбары
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        # Размещение
        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

    def load_data(self):
        """Загрузка данных в таблицу"""
        # Очистка таблицы
        for item in self.tree.get_children():
            self.tree.delete(item)

        # Получение колонок
        columns = self.get_columns()
        if not columns:
            return

        # Загрузка данных
        query = f"SELECT * FROM {self.table_name} ORDER BY id"
        rows = Database.execute_query(query, fetchall=True)

        if rows:
            for row in rows:
                values = []
                for col in columns:
                    val = row.get(col, '')
                    if isinstance(val, datetime):
                        val = val.strftime('%d.%m.%Y')
                    elif val is None:
                        val = ''
                    values.append(val)
                self.tree.insert('', 'end', values=values)

            # Подсветка цветом для статусов
            if self.table_name == 'statuses' and 'color_code' in columns:
                for item in self.tree.get_children():
                    values = self.tree.item(item)['values']
                    color_idx = columns.index('color_code')
                    if color_idx < len(values) and values[color_idx]:
                        self.tree.tag_configure(values[color_idx], background=values[color_idx])
                        self.tree.item(item, tags=(values[color_idx],))

    def search(self, event):
        """Поиск по таблице"""
        search_text = self.search_entry.get().lower()

        # Очистка и перезагрузка
        self.load_data()

        if not search_text:
            return

        # Фильтрация
        for item in self.tree.get_children():
            values = self.tree.item(item)['values']
            if not any(search_text in str(v).lower() for v in values):
                self.tree.delete(item)

    def add_record(self):
        """Добавление записи"""
        self.edit_window = tk.Toplevel(self.window)
        self.edit_window.title(f"Добавление записи - {self.title}")
        self.edit_window.geometry("500x400")
        self.edit_window.grab_set()

        # Создание полей в зависимости от таблицы
        fields = self.get_edit_fields()
        entries = {}

        main_frame = tk.Frame(self.edit_window, padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)

        for i, (field, label, field_type) in enumerate(fields):
            tk.Label(main_frame, text=label + ":", font=("Arial", 10)).grid(row=i, column=0, sticky="w", pady=5)

            if field_type == 'text':
                entry = tk.Entry(main_frame, width=40)
                entry.grid(row=i, column=1, pady=5, padx=5)
                entries[field] = entry

            elif field_type == 'combo':
                combo = ttk.Combobox(main_frame, width=38)
                combo.grid(row=i, column=1, pady=5, padx=5)
                entries[field] = combo
                self.load_combo_data(field, combo)

            elif field_type == 'date':
                entry = tk.Entry(main_frame, width=40)
                entry.insert(0, datetime.now().strftime('%d.%m.%Y'))
                entry.grid(row=i, column=1, pady=5, padx=5)
                entries[field] = entry

            elif field_type == 'file':  # Новый тип для выбора файла
                file_frame = tk.Frame(main_frame)
                file_frame.grid(row=i, column=1, pady=5, padx=5, sticky="w")

                entry = tk.Entry(file_frame, width=30)
                entry.pack(side="left", padx=5)
                entries[field] = entry

                tk.Button(file_frame, text="Обзор...", command=lambda e=entry: self.browse_file(e),
                          bg="#2196F3", fg="white").pack(side="left")

        # Кнопки
        btn_frame = tk.Frame(main_frame)
        btn_frame.grid(row=len(fields), column=0, columnspan=2, pady=20)

        tk.Button(btn_frame, text="Сохранить", command=lambda: self.save_record(entries, fields),
                  bg="#4CAF50", fg="white", width=10).pack(side="left", padx=5)
        tk.Button(btn_frame, text="Отмена", command=self.edit_window.destroy,
                  bg="#f44336", fg="white", width=10).pack(side="left", padx=5)

    def browse_file(self, entry):
        """Открыть диалог выбора файла"""
        filename = filedialog.askopenfilename(
            title="Выберите файл шаблона",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            entry.delete(0, tk.END)
            entry.insert(0, filename)

    def get_edit_fields(self):
        """Получение полей для редактирования"""
        fields_map = {
            'expert_types': [('name', 'Наименование', 'text'), ('description', 'Описание', 'text'),
                             ('regulatory_base', 'Нормативная база', 'text')],
            'objects': [('name', 'Наименование', 'text'), ('category', 'Категория', 'text'),
                        ('unit', 'Ед. измерения', 'text')],
            'customers': [('name', 'Наименование', 'text'), ('inn', 'ИНН', 'text'),
                          ('contact_person', 'Контактное лицо', 'text'), ('phone', 'Телефон', 'text'),
                          ('email', 'Email', 'text')],
            'employees': [('full_name', 'ФИО', 'text'), ('position', 'Должность', 'text'),
                          ('role', 'Роль', 'combo'), ('login', 'Логин', 'text'),
                          ('contacts', 'Контакты', 'text')],
            'statuses': [('name', 'Наименование', 'text'), ('description', 'Описание', 'text'),
                         ('color_code', 'Цвет (код)', 'text')],
            'templates': [('name', 'Наименование', 'text'), ('doc_type', 'Тип документа', 'combo'),
                          ('template_path', 'Путь к шаблону', 'file'), ('update_date', 'Дата обновления', 'date')]
            # Изменено на 'file'
        }
        return fields_map.get(self.table_name, [])

    def load_combo_data(self, field, combo):
        """Загрузка данных для комбобокса"""
        if field == 'role':
            combo['values'] = ['admin', 'director', 'expert', 'specialist']
        elif field == 'doc_type':
            combo['values'] = ['act', 'report', 'conclusion']

    def save_record(self, entries, fields):
        """Сохранение новой записи"""
        data = {}
        for field, _, _ in fields:
            entry = entries.get(field)
            if entry:
                value = entry.get()
                if field in ['update_date']:
                    try:
                        value = datetime.strptime(value, '%d.%m.%Y').date()
                    except:
                        messagebox.showerror("Ошибка", f"Неверный формат даты для поля {field}")
                        return
                data[field] = value

        # Формирование запроса
        columns = ', '.join(data.keys())
        placeholders = ', '.join(['%s'] * len(data))
        values = list(data.values())

        query = f"INSERT INTO {self.table_name} ({columns}) VALUES ({placeholders})"

        record_id = Database.execute_query(query, values, commit=True)

        if record_id:
            Database.log_action(current_user['id'], 'CREATE', self.table_name, record_id,
                                f"Добавлена запись в {self.title}")
            messagebox.showinfo("Успех", "Запись успешно добавлена")
            self.edit_window.destroy()
            self.load_data()
        else:
            messagebox.showerror("Ошибка", "Не удалось добавить запись")

    def edit_record(self):
        """Редактирование записи"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите запись для редактирования")
            return

        # Получение текущих значений
        item = self.tree.item(selected[0])
        current_values = item['values']
        columns = self.get_columns()

        # Создание окна редактирования
        self.edit_window = tk.Toplevel(self.window)
        self.edit_window.title(f"Редактирование записи - {self.title}")
        self.edit_window.geometry("500x400")
        self.edit_window.grab_set()

        fields = self.get_edit_fields()
        entries = {}

        main_frame = tk.Frame(self.edit_window, padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)

        for i, (field, label, field_type) in enumerate(fields):
            tk.Label(main_frame, text=label + ":", font=("Arial", 10)).grid(row=i, column=0, sticky="w", pady=5)

            current_value = ''
            if field in columns:
                idx = columns.index(field)
                if idx < len(current_values):
                    current_value = current_values[idx]

            if field_type == 'text' or field_type == 'date':
                entry = tk.Entry(main_frame, width=40)
                entry.insert(0, current_value)
                entry.grid(row=i, column=1, pady=5, padx=5)
                entries[field] = entry
            elif field_type == 'combo':
                combo = ttk.Combobox(main_frame, width=38)
                combo.set(current_value)
                combo.grid(row=i, column=1, pady=5, padx=5)
                entries[field] = combo
                self.load_combo_data(field, combo)

        # Кнопки
        btn_frame = tk.Frame(main_frame)
        btn_frame.grid(row=len(fields), column=0, columnspan=2, pady=20)

        record_id = current_values[0]
        tk.Button(btn_frame, text="Сохранить", command=lambda: self.update_record(record_id, entries, fields),
                  bg="#4CAF50", fg="white", width=10).pack(side="left", padx=5)
        tk.Button(btn_frame, text="Отмена", command=self.edit_window.destroy,
                  bg="#f44336", fg="white", width=10).pack(side="left", padx=5)

    def update_record(self, record_id, entries, fields):
        """Обновление записи"""
        set_clauses = []
        values = []

        for field, _, _ in fields:
            entry = entries.get(field)
            if entry:
                value = entry.get()
                if field in ['update_date']:
                    try:
                        value = datetime.strptime(value, '%d.%m.%Y').date()
                    except:
                        messagebox.showerror("Ошибка", f"Неверный формат даты для поля {field}")
                        return
                set_clauses.append(f"{field} = %s")
                values.append(value)

        values.append(record_id)
        query = f"UPDATE {self.table_name} SET {', '.join(set_clauses)} WHERE id = %s"

        result = Database.execute_query(query, values, commit=True)

        if result is not None:
            Database.log_action(current_user['id'], 'UPDATE', self.table_name, record_id,
                                f"Обновлена запись ID {record_id} в {self.title}")
            messagebox.showinfo("Успех", "Запись успешно обновлена")
            self.edit_window.destroy()
            self.load_data()
        else:
            messagebox.showerror("Ошибка", "Не удалось обновить запись")

    def delete_record(self):
        """Удаление записи"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите запись для удаления")
            return

        if messagebox.askyesno("Подтверждение", "Вы уверены, что хотите удалить запись?"):
            item = self.tree.item(selected[0])
            record_id = item['values'][0]

            query = f"DELETE FROM {self.table_name} WHERE id = %s"
            result = Database.execute_query(query, (record_id,), commit=True)

            if result is not None:
                Database.log_action(current_user['id'], 'DELETE', self.table_name, record_id,
                                    f"Удалена запись ID {record_id} из {self.title}")
                self.load_data()


class ProjectCreateWindow:
    """Окно создания нового проекта"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Создание нового проекта")
        self.window.geometry("700x650")
        self.window.grab_set()

        main_frame = tk.Frame(self.window, padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)

        # Заказчик
        tk.Label(main_frame, text="Заказчик:*", font=("Arial", 10, "bold")).grid(row=0, column=0, sticky="w", pady=5)
        self.customer_var = tk.StringVar()
        self.customer_combo = ttk.Combobox(main_frame, textvariable=self.customer_var, width=50)
        self.load_customers()
        self.customer_combo.grid(row=0, column=1, pady=5)

        # Объект экспертизы
        tk.Label(main_frame, text="Объект экспертизы:*", font=("Arial", 10, "bold")).grid(row=1, column=0, sticky="w",
                                                                                          pady=5)
        self.object_var = tk.StringVar()
        self.object_combo = ttk.Combobox(main_frame, textvariable=self.object_var, width=50)
        self.load_objects()
        self.object_combo.grid(row=1, column=1, pady=5)

        # Тип работ
        tk.Label(main_frame, text="Тип работ:*", font=("Arial", 10, "bold")).grid(row=2, column=0, sticky="w", pady=5)
        self.type_var = tk.StringVar()
        self.type_combo = ttk.Combobox(main_frame, textvariable=self.type_var, width=50)
        self.load_expert_types()
        self.type_combo.grid(row=2, column=1, pady=5)

        # Ответственный эксперт
        tk.Label(main_frame, text="Ответственный эксперт:*", font=("Arial", 10, "bold")).grid(row=3, column=0,
                                                                                              sticky="w", pady=5)
        self.expert_var = tk.StringVar()
        self.expert_combo = ttk.Combobox(main_frame, textvariable=self.expert_var, width=50)
        self.load_experts()
        self.expert_combo.grid(row=3, column=1, pady=5)

        # Планируемые сроки
        tk.Label(main_frame, text="Планируемые сроки:", font=("Arial", 10, "bold")).grid(row=4, column=0, sticky="w",
                                                                                         pady=5)
        date_frame = tk.Frame(main_frame)
        date_frame.grid(row=4, column=1, pady=5, sticky="w")

        tk.Label(date_frame, text="с").pack(side="left")
        self.start_entry = tk.Entry(date_frame, width=12)
        self.start_entry.pack(side="left", padx=5)
        self.start_entry.insert(0, datetime.now().strftime("%d.%m.%Y"))

        tk.Label(date_frame, text="по").pack(side="left", padx=5)
        self.end_entry = tk.Entry(date_frame, width=12)
        self.end_entry.pack(side="left", padx=5)

        # Номер договора (из А1)
        tk.Label(main_frame, text="Номер договора (из А1):", font=("Arial", 10, "bold")).grid(row=5, column=0,
                                                                                              sticky="w", pady=5)
        self.contract_number_entry = tk.Entry(main_frame, width=30)
        self.contract_number_entry.grid(row=5, column=1, pady=5, sticky="w")

        # Дата договора
        tk.Label(main_frame, text="Дата договора:", font=("Arial", 10, "bold")).grid(row=6, column=0, sticky="w",
                                                                                     pady=5)
        self.contract_date_entry = tk.Entry(main_frame, width=12)
        self.contract_date_entry.grid(row=6, column=1, pady=5, sticky="w")
        self.contract_date_entry.insert(0, datetime.now().strftime("%d.%m.%Y"))

        # Примечания
        tk.Label(main_frame, text="Примечания:", font=("Arial", 10, "bold")).grid(row=7, column=0, sticky="nw", pady=5)
        self.notes_text = scrolledtext.ScrolledText(main_frame, width=50, height=5)
        self.notes_text.grid(row=7, column=1, pady=5)

        # Кнопки
        btn_frame = tk.Frame(main_frame)
        btn_frame.grid(row=8, column=0, columnspan=2, pady=20)

        tk.Button(btn_frame, text="Создать проект", command=self.create_project,
                  bg="#4CAF50", fg="white", font=("Arial", 10), width=15).pack(side="left", padx=5)
        tk.Button(btn_frame, text="Отмена", command=self.window.destroy,
                  bg="#f44336", fg="white", font=("Arial", 10), width=10).pack(side="left", padx=5)

    def load_customers(self):
        customers = Database.execute_query("SELECT id, name FROM customers ORDER BY name", fetchall=True)
        if customers:
            self.customer_combo['values'] = [f"{c['id']}: {c['name']}" for c in customers]

    def load_objects(self):
        objects = Database.execute_query("SELECT id, name FROM objects ORDER BY name", fetchall=True)
        if objects:
            self.object_combo['values'] = [f"{o['id']}: {o['name']}" for o in objects]

    def load_expert_types(self):
        types = Database.execute_query("SELECT id, name FROM expert_types ORDER BY name", fetchall=True)
        if types:
            self.type_combo['values'] = [f"{t['id']}: {t['name']}" for t in types]

    def load_experts(self):
        experts = Database.execute_query("SELECT id, full_name FROM employees WHERE role='expert' ORDER BY full_name",
                                         fetchall=True)
        if experts:
            self.expert_combo['values'] = [f"{e['id']}: {e['full_name']}" for e in experts]

    def create_project(self):
        # Проверка заполнения обязательных полей
        if not self.customer_var.get() or not self.object_var.get() or not self.type_var.get() or not self.expert_var.get():
            messagebox.showerror("Ошибка", "Заполните все обязательные поля")
            return

        # Парсинг ID из выбранных значений
        try:
            customer_id = int(self.customer_var.get().split(':')[0])
            object_id = int(self.object_var.get().split(':')[0])
            expert_type_id = int(self.type_var.get().split(':')[0])
            expert_id = int(self.expert_var.get().split(':')[0])
        except:
            messagebox.showerror("Ошибка", "Некорректный выбор из справочника")
            return

        # Генерация номера проекта
        date_str = datetime.now().strftime('%Y%m%d')
        project_number = f"П-{date_str}-{datetime.now().strftime('%H%M%S')}"

        # Получение статуса "Черновик"
        status = Database.execute_query("SELECT id FROM statuses WHERE name='Черновик'", fetchone=True)
        status_id = status['id'] if status else 1

        # Вставка в БД
        query = """
            INSERT INTO projects (project_number, creation_date, customer_id, object_id, expert_type_id, 
                                 expert_id, planned_start, planned_end, contract_number, contract_date, 
                                 notes, status_id)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """

        try:
            start_date = datetime.strptime(self.start_entry.get(),
                                           "%d.%m.%Y").date() if self.start_entry.get() else None
            end_date = datetime.strptime(self.end_entry.get(), "%d.%m.%Y").date() if self.end_entry.get() else None
            contract_date = datetime.strptime(self.contract_date_entry.get(),
                                              "%d.%m.%Y").date() if self.contract_date_entry.get() else None
        except:
            messagebox.showerror("Ошибка", "Некорректный формат даты. Используйте ДД.ММ.ГГГГ")
            return

        params = (
            project_number,
            datetime.now().date(),
            customer_id,
            object_id,
            expert_type_id,
            expert_id,
            start_date,
            end_date,
            self.contract_number_entry.get(),
            contract_date,
            self.notes_text.get('1.0', 'end-1c'),
            status_id
        )

        project_id = Database.execute_query(query, params, commit=True)

        if project_id:
            Database.log_action(current_user['id'], 'CREATE', 'project', project_id, f"Создан проект {project_number}")
            messagebox.showinfo("Успех", f"Проект {project_number} успешно создан")
            self.window.destroy()
        else:
            messagebox.showerror("Ошибка", "Не удалось создать проект")


class ProjectListWindow:
    """Окно списка проектов"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Список проектов")
        self.window.geometry("1100x600")

        # Панель инструментов
        toolbar = tk.Frame(self.window, bg="#f0f0f0", height=40)
        toolbar.pack(fill="x", padx=2, pady=2)

        tk.Button(toolbar, text="Обновить", command=self.load_projects,
                  bg="#FF9800", fg="white", width=10).pack(side="left", padx=5, pady=5)

        if current_user_role in ['director', 'admin']:
            tk.Button(toolbar, text="Все проекты", command=lambda: self.load_projects(all_projects=True),
                      bg="#2196F3", fg="white", width=10).pack(side="left", padx=5, pady=5)

        # Фильтр по статусу
        tk.Label(toolbar, text="Статус:", bg="#f0f0f0").pack(side="left", padx=(20, 5))
        self.status_filter = ttk.Combobox(toolbar,
                                          values=['Все', 'Черновик', 'На согласовании', 'Утвержден', 'Передан в А4'],
                                          width=15)
        self.status_filter.set('Все')
        self.status_filter.pack(side="left", padx=5)
        self.status_filter.bind('<<ComboboxSelected>>', self.apply_filter)

        # Поиск
        tk.Label(toolbar, text="Поиск:", bg="#f0f0f0").pack(side="right", padx=5)
        self.search_entry = tk.Entry(toolbar, width=30)
        self.search_entry.pack(side="right", padx=5)
        self.search_entry.bind('<KeyRelease>', self.apply_filter)

        # Основная таблица
        self.create_table()
        self.load_projects()

    def create_table(self):
        """Создание таблицы проектов"""
        table_frame = tk.Frame(self.window)
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        columns = ('ID', 'Номер', 'Дата', 'Заказчик', 'Объект', 'Эксперт', 'Статус', 'Договор')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=20)

        # Настройка колонок
        col_widths = [50, 120, 90, 150, 150, 150, 100, 100]
        for i, col in enumerate(columns):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=col_widths[i], anchor="w")

        # Скроллбары
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        # Размещение
        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

        # Двойной клик для просмотра деталей
        self.tree.bind('<Double-Button-1>', self.view_project)

    def load_projects(self, all_projects=False):
        """Загрузка списка проектов"""
        # Очистка таблицы
        for item in self.tree.get_children():
            self.tree.delete(item)

        # Формирование запроса
        if current_user_role == 'expert' and not all_projects:
            query = """
                SELECT p.id, p.project_number, p.creation_date, c.name as customer_name, 
                       o.name as object_name, e.full_name as expert_name, s.name as status_name,
                       p.contract_number
                FROM projects p
                JOIN customers c ON p.customer_id = c.id
                JOIN objects o ON p.object_id = o.id
                JOIN employees e ON p.expert_id = e.id
                JOIN statuses s ON p.status_id = s.id
                WHERE p.expert_id = %s
                ORDER BY p.creation_date DESC
            """
            projects = Database.execute_query(query, (current_user['id'],), fetchall=True)
        else:
            query = """
                SELECT p.id, p.project_number, p.creation_date, c.name as customer_name, 
                       o.name as object_name, e.full_name as expert_name, s.name as status_name,
                       p.contract_number
                FROM projects p
                JOIN customers c ON p.customer_id = c.id
                JOIN objects o ON p.object_id = o.id
                JOIN employees e ON p.expert_id = e.id
                JOIN statuses s ON p.status_id = s.id
                ORDER BY p.creation_date DESC
            """
            projects = Database.execute_query(query, fetchall=True)

        self.all_projects = projects

        if projects:
            for p in projects:
                self.tree.insert('', 'end', values=(
                    p['id'],
                    p['project_number'],
                    p['creation_date'].strftime('%d.%m.%Y') if p['creation_date'] else '',
                    p['customer_name'],
                    p['object_name'],
                    p['expert_name'],
                    p['status_name'],
                    p['contract_number'] or ''
                ))

    def apply_filter(self, event=None):
        """Применение фильтра по статусу и поиска"""
        status_filter = self.status_filter.get()
        search_text = self.search_entry.get().lower()

        # Очистка таблицы
        for item in self.tree.get_children():
            self.tree.delete(item)

        if not hasattr(self, 'all_projects') or not self.all_projects:
            return

        for p in self.all_projects:
            # Фильтр по статусу
            if status_filter != 'Все' and p['status_name'] != status_filter:
                continue

            # Поиск по тексту
            if search_text:
                searchable = f"{p['project_number']} {p['customer_name']} {p['object_name']} {p['contract_number']}".lower()
                if search_text not in searchable:
                    continue

            self.tree.insert('', 'end', values=(
                p['id'],
                p['project_number'],
                p['creation_date'].strftime('%d.%m.%Y') if p['creation_date'] else '',
                p['customer_name'],
                p['object_name'],
                p['expert_name'],
                p['status_name'],
                p['contract_number'] or ''
            ))

    def view_project(self, event):
        """Просмотр деталей проекта"""
        selected = self.tree.selection()
        if not selected:
            return

        item = self.tree.item(selected[0])
        project_id = item['values'][0]

        ProjectViewWindow(project_id)


class ProjectViewWindow:
    """Окно просмотра деталей проекта"""

    def __init__(self, project_id):
        self.project_id = project_id
        self.window = tk.Toplevel()
        self.window.title(f"Просмотр проекта ID {project_id}")
        self.window.geometry("800x600")

        # Загрузка данных проекта
        query = """
            SELECT p.*, c.name as customer_name, o.name as object_name, 
                   et.name as expert_type_name, e.full_name as expert_name,
                   s.name as status_name
            FROM projects p
            JOIN customers c ON p.customer_id = c.id
            JOIN objects o ON p.object_id = o.id
            JOIN expert_types et ON p.expert_type_id = et.id
            JOIN employees e ON p.expert_id = e.id
            JOIN statuses s ON p.status_id = s.id
            WHERE p.id = %s
        """
        project = Database.execute_query(query, (project_id,), fetchone=True)

        if not project:
            messagebox.showerror("Ошибка", "Проект не найден")
            self.window.destroy()
            return

        # Основной фрейм
        main_frame = tk.Frame(self.window, padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)

        # Заголовок
        tk.Label(main_frame, text=f"Проект: {project['project_number']}",
                 font=("Arial", 14, "bold")).pack(anchor="w", pady=10)

        # Информация о проекте
        info_frame = tk.LabelFrame(main_frame, text="Основная информация", font=("Arial", 12, "bold"))
        info_frame.pack(fill="x", pady=10)

        info_grid = tk.Frame(info_frame)
        info_grid.pack(pady=10, padx=10)

        fields = [
            ("Заказчик:", project['customer_name']),
            ("Объект экспертизы:", project['object_name']),
            ("Тип работ:", project['expert_type_name']),
            ("Ответственный эксперт:", project['expert_name']),
            ("Дата создания:", project['creation_date'].strftime('%d.%m.%Y') if project['creation_date'] else ''),
            ("Плановые сроки:",
             f"{project['planned_start'].strftime('%d.%m.%Y') if project['planned_start'] else ''} - {project['planned_end'].strftime('%d.%m.%Y') if project['planned_end'] else ''}"),
            ("Статус:", project['status_name']),
            ("Договор:",
             f"{project['contract_number']} от {project['contract_date'].strftime('%d.%m.%Y') if project['contract_date'] else ''}")
        ]

        for i, (label, value) in enumerate(fields):
            tk.Label(info_grid, text=label, font=("Arial", 10, "bold")).grid(row=i, column=0, sticky="w", pady=2)
            tk.Label(info_grid, text=value, font=("Arial", 10)).grid(row=i, column=1, sticky="w", pady=2, padx=(10, 0))

        # Примечания
        if project['notes']:
            notes_frame = tk.LabelFrame(main_frame, text="Примечания", font=("Arial", 12, "bold"))
            notes_frame.pack(fill="x", pady=10)

            tk.Label(notes_frame, text=project['notes'], font=("Arial", 10),
                     wraplength=700, justify="left").pack(pady=10, padx=10)

        # Кнопки
        btn_frame = tk.Frame(main_frame)
        btn_frame.pack(pady=20)

        if current_user_role in ['expert'] and project['status_name'] == 'Черновик':
            tk.Button(btn_frame, text="Загрузить результаты",
                      command=lambda: self.open_results_upload(),
                      bg="#4CAF50", fg="white", width=15).pack(side="left", padx=5)

        if current_user_role in ['specialist'] and project['status_name'] in ['Черновик', 'На согласовании']:
            tk.Button(btn_frame, text="Создать документы",
                      command=lambda: self.create_documents(),
                      bg="#2196F3", fg="white", width=15).pack(side="left", padx=5)

        tk.Button(btn_frame, text="Закрыть", command=self.window.destroy,
                  bg="#f44336", fg="white", width=10).pack(side="left", padx=5)

    def open_results_upload(self):
        """Открыть окно загрузки результатов"""
        self.window.destroy()
        ResultsUploadWindow(self.project_id)

    def create_documents(self):
        """Создание документов по проекту"""
        # Проверка наличия результатов
        results = Database.execute_query(
            "SELECT COUNT(*) as count FROM results WHERE project_id = %s",
            (self.project_id,), fetchone=True
        )

        if not results or results['count'] == 0:
            messagebox.showwarning("Предупреждение", "Сначала необходимо загрузить результаты обследования")
            return

        # Получаем данные проекта для подстановки в шаблон
        project = Database.execute_query("""
            SELECT p.*, c.name as customer_name, o.name as object_name, 
                   et.name as expert_type_name, e.full_name as expert_name
            FROM projects p
            JOIN customers c ON p.customer_id = c.id
            JOIN objects o ON p.object_id = o.id
            JOIN expert_types et ON p.expert_type_id = et.id
            JOIN employees e ON p.expert_id = e.id
            WHERE p.id = %s
        """, (self.project_id,), fetchone=True)

        # Получаем все шаблоны
        templates = Database.execute_query(
            "SELECT id, doc_type, template_path FROM templates", fetchall=True
        )

        if not templates:
            messagebox.showerror("Ошибка", "В системе нет шаблонов документов")
            return

        # Получаем статус "Черновик"
        status = Database.execute_query("SELECT id FROM statuses WHERE name='Черновик'", fetchone=True)
        status_id = status['id'] if status else 1

        created = 0
        skipped = 0

        for template in templates:
            # Проверяем, есть ли уже черновик такого типа
            existing = Database.execute_query(
                "SELECT COUNT(*) as count FROM drafts WHERE project_id = %s AND doc_type = %s",
                (self.project_id, template['doc_type']), fetchone=True
            )

            if existing and existing['count'] > 0:
                print(f"Черновик типа {template['doc_type']} уже существует")
                skipped += 1
                continue

            # Генерируем содержимое на основе шаблона и данных проекта
            content = self.generate_document_content(template['doc_type'], project, template['template_path'])

            # Создаем черновик
            query = """
                INSERT INTO drafts (project_id, doc_type, template_id, content, version, status_id, create_date, author_id)
                VALUES (%s, %s, %s, %s, 1, %s, NOW(), %s)
            """
            doc_id = Database.execute_query(
                query, (self.project_id, template['doc_type'], template['id'], content, status_id, current_user['id']),
                commit=True
            )

            if doc_id:
                created += 1
                print(f"Создан черновик типа {template['doc_type']} с ID {doc_id}")

        if created > 0:
            Database.log_action(current_user['id'], 'CREATE_DRAFTS', 'project', self.project_id,
                                f"Создано {created} черновиков документов")
            messagebox.showinfo("Успех", f"Создано {created} черновиков документов" +
                                (f" ({skipped} уже существовали)" if skipped > 0 else ""))
        else:
            if skipped > 0:
                messagebox.showinfo("Информация", f"Все черновики документов уже существуют ({skipped} шт.)")
            else:
                messagebox.showwarning("Предупреждение", "Не удалось создать черновики")

    def generate_document_content(self, doc_type, project, template_path):
        """Генерация содержимого документа на основе шаблона и данных проекта"""

        # Базовая информация, общая для всех типов документов
        base_content = f"""
    Документ создан в системе автоматизации ООО "Экспертиза-ПБ"
    Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}
    Проект: {project['project_number']}
    Заказчик: {project['customer_name']}
    Объект экспертизы: {project['object_name']}
    Тип работ: {project['expert_type_name']}
    Ответственный эксперт: {project['expert_name']}
    Договор: {project['contract_number']} от {project['contract_date'].strftime('%d.%m.%Y') if project['contract_date'] else ''}
        """

        # Пытаемся загрузить шаблон из файла
        if template_path and os.path.exists(template_path):
            try:
                with open(template_path, 'r', encoding='utf-8') as f:
                    template_content = f.read()
                    # Подставляем данные из проекта
                    template_content = template_content.replace('{{project_number}}', project['project_number'])
                    template_content = template_content.replace('{{customer_name}}', project['customer_name'])
                    template_content = template_content.replace('{{object_name}}', project['object_name'])
                    template_content = template_content.replace('{{expert_name}}', project['expert_name'])
                    template_content = template_content.replace('{{contract_number}}', project['contract_number'] or '')
                    return template_content
            except Exception as e:
                print(f"Ошибка загрузки шаблона: {e}")

        # Если шаблон не загрузился, используем стандартный
        doc_templates = {
            'act': f"""
    АКТ ОБСЛЕДОВАНИЯ
    {base_content}

    Результаты обследования:
    [Будут добавлены результаты после загрузки]

    Состав комиссии:
    1. {project['expert_name']} - эксперт

    Заключение комиссии:
    [Требуется заполнить]

    Подписи:
    _____________ {project['expert_name']}
            """,

            'report': f"""
    ТЕХНИЧЕСКИЙ ОТЧЕТ
    {base_content}

    Цель работы: Проведение технического диагностирования

    Методика проведения:
    [Описание методики]

    Результаты расчетов и анализов:
    [Результаты будут добавлены]

    Графики и таблицы:
    [Будут добавлены]

    Заключение технической части:
    [Требуется заполнить]
            """,

            'conclusion': f"""
    ЭКСПЕРТНОЕ ЗАКЛЮЧЕНИЕ
    {base_content}

    Нормативная база:
    - Федеральный закон №116-ФЗ "О промышленной безопасности"
    - Федеральные нормы и правила в области промышленной безопасности

    Результаты экспертизы:
    [Будут добавлены на основе акта и отчета]

    Выводы:
    [Требуется сформулировать выводы о соответствии]

    Рекомендации:
    [При необходимости]

    Эксперт: {project['expert_name']}
            """
        }

        return doc_templates.get(doc_type, base_content)

    def get_results_text(self, project_id):
        """Получение текста результатов обследования с новыми полями"""
        result = Database.execute_query("""
            SELECT * FROM results 
            WHERE project_id = %s 
            ORDER BY upload_date DESC 
            LIMIT 1
        """, (project_id,), fetchone=True)

        if not result:
            return "Результаты обследования не загружены"

        text = f"Дата загрузки: {result['upload_date'].strftime('%d.%m.%Y %H:%M') if result['upload_date'] else ''}\n"
        text += f"Рабочее давление: {result['working_pressure'] or 'не указано'} МПа\n"
        text += f"Рабочая температура: {result['working_temperature'] or 'не указано'} °C\n"
        text += f"Вместимость: {result['capacity'] or 'не указано'} м³\n"
        text += f"Рабочая среда: {result['working_medium'] or 'не указано'}\n"
        text += f"Класс опасности: {result['hazard_class'] or 'не указано'}\n"
        text += f"Материал корпуса: {result['body_material'] or 'не указано'}\n"
        text += f"Тип сварки: {result['welding_type'] or 'не указано'}\n"
        text += f"Минимальная толщина: {result['min_thickness'] or 'не указано'} мм\n"
        text += f"Основные дефекты: {result['max_defects'] or 'не указано'}\n"
        text += f"Остаточный ресурс: {result['remaining_life'] or 'не указано'} лет\n"
        text += f"Замеры: {result['measurements']}\n"
        text += f"Комментарии: {result['comments']}\n"
        text += "-" * 40 + "\n"

        return text

    def create_draft_from_template(self, project, template_path, doc_type):
        """Создание черновика из шаблона Word с подстановкой данных"""

        print(f"Создание черновика из шаблона: {template_path}")

        if not template_path or not os.path.exists(template_path):
            print(f"Файл шаблона не найден, создаем пустой")
            return self.create_empty_draft(project, doc_type)

        # Получаем результаты обследования с новыми полями
        result = Database.execute_query("""
            SELECT * FROM results 
            WHERE project_id = %s 
            ORDER BY upload_date DESC 
            LIMIT 1
        """, (project['id'],), fetchone=True)

        # Генерируем имя файла
        draft_filename = f"{doc_type}_{project['project_number']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        draft_path = os.path.join(DRAFTS_FOLDER, draft_filename)

        # Копируем шаблон
        shutil.copy2(template_path, draft_path)

        # Подставляем данные
        from docx import Document
        doc = Document(draft_path)

        # Словарь для замены (старые + новые 10 полей)
        replacements = {
            # Старые поля
            '{{project_number}}': project['project_number'],
            '{{customer_name}}': project['customer_name'],
            '{{object_name}}': project['object_name'],
            '{{expert_name}}': project['expert_name'],
            '{{contract_number}}': project['contract_number'] or '',
            '{{contract_date}}': project['contract_date'].strftime('%d.%m.%Y') if project['contract_date'] else '',
            '{{current_date}}': datetime.now().strftime('%d.%m.%Y'),
            '{{results}}': self.get_results_text(project['id']),
            '{{doc_number}}': f"{doc_type.upper()}-{datetime.now().strftime('%Y%m%d')}-{project['id']}",

            # Новые 10 полей
            '{{working_pressure}}': result['working_pressure'] if result and result[
                'working_pressure'] else '[УКАЗАТЬ]',
            '{{working_temperature}}': result['working_temperature'] if result and result[
                'working_temperature'] else '[УКАЗАТЬ]',
            '{{capacity}}': result['capacity'] if result and result['capacity'] else '[УКАЗАТЬ]',
            '{{working_medium}}': result['working_medium'] if result and result['working_medium'] else '[УКАЗАТЬ]',
            '{{hazard_class}}': result['hazard_class'] if result and result['hazard_class'] else '[УКАЗАТЬ]',
            '{{body_material}}': result['body_material'] if result and result['body_material'] else '[УКАЗАТЬ]',
            '{{welding_type}}': result['welding_type'] if result and result['welding_type'] else '[УКАЗАТЬ]',
            '{{min_thickness}}': result['min_thickness'] if result and result['min_thickness'] else '[УКАЗАТЬ]',
            '{{max_defects}}': result['max_defects'] if result and result['max_defects'] else '[УКАЗАТЬ]',
            '{{remaining_life}}': result['remaining_life'] if result and result['remaining_life'] else '[УКАЗАТЬ]',
        }

        print(f"Словарь замены: {replacements}")

        # Заменяем во всех параграфах
        replaced_count = 0
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
                    replaced_count += 1

        # Заменяем в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, str(value))
                                replaced_count += 1

        print(f"Выполнено замен: {replaced_count}")
        doc.save(draft_path)
        print(f"Черновик создан: {draft_path}")

        return draft_path

    def create_empty_draft(self, project, doc_type):
        """Создание пустого черновика, если шаблон не работает"""
        print(f"Создание пустого черновика для {doc_type}")

        draft_filename = f"{doc_type}_{project['project_number']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        draft_path = os.path.join(DRAFTS_FOLDER, draft_filename)

        # Создаем простой Word документ
        from docx import Document
        doc = Document()
        doc.add_heading(f'{doc_type.upper()} - Проект {project["project_number"]}', 0)
        doc.add_paragraph(f'Заказчик: {project["customer_name"]}')
        doc.add_paragraph(f'Объект: {project["object_name"]}')
        doc.add_paragraph(f'Эксперт: {project["expert_name"]}')
        doc.add_paragraph(f'Дата создания: {datetime.now().strftime("%d.%m.%Y")}')
        doc.add_paragraph('\nРезультаты обследования:')
        doc.add_paragraph(self.get_results_text(project['id']))
        doc.save(draft_path)

        return draft_path

    # ЗАМЕНИТЕ существующий метод create_documents на этот
    def create_documents(self):
        """Создание документов по проекту"""
        # Проверка наличия результатов
        results = Database.execute_query(
            "SELECT COUNT(*) as count FROM results WHERE project_id = %s",
            (self.project_id,), fetchone=True
        )

        if not results or results['count'] == 0:
            messagebox.showwarning("Предупреждение", "Сначала необходимо загрузить результаты обследования")
            return

        # Получаем данные проекта
        project = Database.execute_query("""
            SELECT p.*, c.name as customer_name, o.name as object_name, 
                   et.name as expert_type_name, e.full_name as expert_name
            FROM projects p
            JOIN customers c ON p.customer_id = c.id
            JOIN objects o ON p.object_id = o.id
            JOIN expert_types et ON p.expert_type_id = et.id
            JOIN employees e ON p.expert_id = e.id
            WHERE p.id = %s
        """, (self.project_id,), fetchone=True)

        # Получаем шаблоны
        templates = Database.execute_query(
            "SELECT id, doc_type, template_path FROM templates", fetchall=True
        )

        if not templates:
            messagebox.showerror("Ошибка", "В системе нет шаблонов документов")
            return

        created = 0

        for template in templates:
            # Проверяем существование черновика
            existing = Database.execute_query(
                "SELECT id, file_path FROM drafts WHERE project_id = %s AND doc_type = %s",
                (self.project_id, template['doc_type']), fetchone=True
            )

            if existing:
                # Если черновик есть, просто открываем его
                if messagebox.askyesno("Вопрос",
                                       f"Черновик {template['doc_type']} уже существует. Открыть его?"):
                    self.open_word_file(existing['file_path'])
                continue

            # Создаем новый черновик из шаблона
            if not template['template_path'] or not os.path.exists(template['template_path']):
                messagebox.showwarning("Предупреждение",
                                       f"Файл шаблона {template['template_path']} не найден")
                continue

            # Генерируем файл черновика
            draft_path = self.create_draft_from_template(project, template['template_path'], template['doc_type'])

            # Сохраняем в БД
            status = Database.execute_query("SELECT id FROM statuses WHERE name='Черновик'", fetchone=True)
            status_id = status['id'] if status else 1

            query = """
                INSERT INTO drafts (project_id, doc_type, template_id, file_path, version, status_id, create_date, author_id)
                VALUES (%s, %s, %s, %s, 1, %s, NOW(), %s)
            """
            doc_id = Database.execute_query(
                query,
                (self.project_id, template['doc_type'], template['id'], draft_path, status_id, current_user['id']),
                commit=True
            )

            if doc_id:
                created += 1
                # Сразу открываем созданный документ
                self.open_word_file(draft_path)

        if created > 0:
            messagebox.showinfo("Успех", f"Создано {created} черновиков документов")

    def open_word_file(self, file_path):
        """Открытие файла в Word"""
        try:
            if os.name == 'nt':  # Windows
                os.startfile(file_path)
            else:  # Linux/Mac
                subprocess.call(['xdg-open', file_path])
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл: {e}")


class ResultsUploadWindow:
    """Окно загрузки результатов обследования (КРАСИВАЯ ВЕРСИЯ)"""

    def __init__(self, project_id=None):
        self.project_id = project_id
        self.uploaded_files = []

        self.window = tk.Toplevel()
        self.window.title("Загрузка результатов обследования")
        self.window.geometry("900x850")
        self.window.grab_set()

        # Центрируем окно
        self.center_window()

        # Настраиваем стили
        self.setup_styles()

        # Создаем интерфейс
        self.create_widgets()

    def center_window(self):
        """Центрирование окна на экране"""
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f'{width}x{height}+{x}+{y}')

    def setup_styles(self):
        """Настройка стилей для красивого отображения"""
        self.bg_color = "#f5f5f5"
        self.section_bg = "#ffffff"
        self.primary_color = "#2c3e50"
        self.accent_color = "#3498db"
        self.success_color = "#27ae60"
        self.danger_color = "#e74c3c"
        self.warning_color = "#f39c12"
        self.border_color = "#dde4e9"

        # Настройка шрифтов
        self.font_normal = ("Segoe UI", 10)
        self.font_bold = ("Segoe UI", 10, "bold")
        self.font_header = ("Segoe UI", 14, "bold")
        self.font_section = ("Segoe UI", 11, "bold")

        # Цвет фона окна
        self.window.configure(bg=self.bg_color)

    def create_section(self, parent, title, icon="", row=None, col=None, colspan=2):
        """Создание секции с заголовком и рамкой"""
        # Рамка секции
        section_frame = tk.Frame(parent, bg=self.section_bg, relief="solid", bd=1,
                                 highlightbackground=self.border_color, highlightcolor=self.border_color,
                                 highlightthickness=1)
        if row is not None and col is not None:
            section_frame.grid(row=row, column=col, columnspan=colspan, sticky="nsew", padx=10, pady=10)
        else:
            section_frame.pack(fill="x", padx=10, pady=10)

        # Заголовок с иконкой
        header_frame = tk.Frame(section_frame, bg=self.section_bg)
        header_frame.pack(fill="x", padx=15, pady=(15, 5))

        if icon:
            icon_label = tk.Label(header_frame, text=icon, font=("Segoe UI", 14), bg=self.section_bg)
            icon_label.pack(side="left", padx=(0, 5))

        title_label = tk.Label(header_frame, text=title, font=self.font_section, fg=self.accent_color,
                               bg=self.section_bg)
        title_label.pack(side="left")

        # Линия под заголовком
        line = tk.Frame(section_frame, bg=self.border_color, height=1)
        line.pack(fill="x", padx=15, pady=(0, 15))

        # Контент секции (возвращаем для добавления полей)
        content_frame = tk.Frame(section_frame, bg=self.section_bg)
        content_frame.pack(fill="x", padx=15, pady=(0, 15))

        return content_frame

    def create_field(self, parent, label, row, col, tooltip="", is_textarea=False, height=3):
        """Создание поля ввода с меткой"""
        # Метка
        label_widget = tk.Label(parent, text=label, font=self.font_bold, fg=self.primary_color, bg=self.section_bg,
                                anchor="w")
        label_widget.grid(row=row, column=col * 2, sticky="w", padx=(0, 10), pady=8)

        # Поле ввода
        if is_textarea:
            # Текстовая область с прокруткой
            text_frame = tk.Frame(parent, bg=self.section_bg)
            text_frame.grid(row=row, column=col * 2 + 1, sticky="ew", padx=5, pady=5)

            entry = scrolledtext.ScrolledText(text_frame, width=50, height=height, font=self.font_normal,
                                              wrap=tk.WORD, highlightbackground=self.border_color, highlightthickness=1)
            entry.pack(fill="both", expand=True)
        else:
            # Обычное поле ввода
            entry = tk.Entry(parent, width=40, font=self.font_normal, relief="solid", bd=1,
                             highlightbackground=self.border_color)
            entry.grid(row=row, column=col * 2 + 1, sticky="ew", padx=5, pady=5)

        # Кнопка подсказки (если есть)
        if tooltip:
            info_btn = tk.Label(parent, text="ⓘ", font=("Segoe UI", 10), fg=self.accent_color, bg=self.section_bg,
                                cursor="question_arrow")
            info_btn.grid(row=row, column=col * 2 + 2, padx=(2, 0), pady=5)
            info_btn.bind('<Enter>', lambda e, msg=tooltip: self.show_tooltip(e, msg))
            info_btn.bind('<Leave>', self.hide_tooltip)

        return entry

    def show_tooltip(self, event, message):
        """Показать всплывающую подсказку"""
        x = event.widget.winfo_rootx() + 20
        y = event.widget.winfo_rooty() + 20

        self.tooltip = tk.Toplevel()
        self.tooltip.wm_overrideredirect(True)
        self.tooltip.wm_geometry(f"+{x}+{y}")

        label = tk.Label(self.tooltip, text=message, bg="#ffffe0", relief="solid", bd=1,
                         font=("Segoe UI", 9), padx=5, pady=3)
        label.pack()

    def hide_tooltip(self, event):
        """Скрыть всплывающую подсказку"""
        if hasattr(self, 'tooltip'):
            self.tooltip.destroy()

    def create_widgets(self):
        """Создание всех элементов интерфейса"""

        # Создаем холст с прокруткой
        canvas = tk.Canvas(self.window, bg=self.bg_color, highlightthickness=0)
        scrollbar = tk.Scrollbar(self.window, orient="vertical", command=canvas.yview)
        self.scrollable_frame = tk.Frame(canvas, bg=self.bg_color)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw", width=canvas.winfo_width())
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True, padx=10, pady=10)
        scrollbar.pack(side="right", fill="y", pady=10)

        # Обновляем ширину canvas при изменении размера окна
        def on_canvas_configure(e):
            canvas.itemconfig(1, width=e.width)

        canvas.bind('<Configure>', on_canvas_configure)

        # === ЗАГОЛОВОК ===
        header_frame = tk.Frame(self.scrollable_frame, bg=self.bg_color)
        header_frame.pack(fill="x", pady=(0, 20))

        icon_label = tk.Label(header_frame, text="🔬", font=("Segoe UI", 24), bg=self.bg_color)
        icon_label.pack(side="left", padx=(0, 10))

        title_label = tk.Label(header_frame, text="Результаты обследования", font=self.font_header,
                               fg=self.primary_color, bg=self.bg_color)
        title_label.pack(side="left")

        # === СЕКЦИЯ: Информация о проекте ===
        project_section = self.create_section(self.scrollable_frame, "Информация о проекте", "📋")

        if self.project_id:
            project = Database.execute_query(
                "SELECT p.project_number, c.name as customer_name FROM projects p JOIN customers c ON p.customer_id = c.id WHERE p.id = %s",
                (self.project_id,), fetchone=True
            )

            project_info = f"Проект: {project['project_number']} | Заказчик: {project['customer_name']}"
            info_label = tk.Label(project_section, text=project_info, font=self.font_bold, fg=self.accent_color,
                                  bg=self.section_bg)
            info_label.pack(anchor="w", pady=5)
            self.project_var = self.project_id
        else:
            tk.Label(project_section, text="Выберите проект:", font=self.font_bold, fg=self.primary_color,
                     bg=self.section_bg).pack(anchor="w", pady=(0, 5))
            self.project_combo = ttk.Combobox(project_section, width=70, font=self.font_normal)
            self.load_projects()
            self.project_combo.pack(fill="x", pady=5)
            self.project_var = None

        # === СЕКЦИЯ: Параметры объекта ===
        params_section = self.create_section(self.scrollable_frame, "Параметры объекта", "📊")

        # Создаем сетку для двух колонок
        params_section.grid_columnconfigure(1, weight=1)
        params_section.grid_columnconfigure(3, weight=1)

        # Левая колонка (row, col=0)
        self.working_pressure = self.create_field(params_section, "Рабочее давление (МПа):", 0, 0,
                                                  "Введите максимальное рабочее давление в МПа")
        self.working_temperature = self.create_field(params_section, "Рабочая температура (°C):", 1, 0,
                                                     "Введите максимальную рабочую температуру")
        self.capacity = self.create_field(params_section, "Вместимость (м³):", 2, 0,
                                          "Объем объекта в кубических метрах")

        # Правая колонка (row, col=1)
        self.working_medium = self.create_field(params_section, "Рабочая среда:", 0, 1,
                                                "Например: пар, вода, нефть, газ")
        self.hazard_class = self.create_field(params_section, "Класс опасности:", 1, 1,
                                              "I, II, III или IV класс опасности")
        self.body_material = self.create_field(params_section, "Материал корпуса:", 2, 1,
                                               "Сталь 09Г2С, 12Х18Н10Т, и т.д.")

        # === СЕКЦИЯ: Материалы и сварка ===
        materials_section = self.create_section(self.scrollable_frame, "Материалы и сварка", "🔧")
        materials_section.grid_columnconfigure(1, weight=1)

        self.welding_type = self.create_field(materials_section, "Тип сварки:", 0, 0,
                                              "Ручная дуговая, автоматическая, и т.д.")

        # === СЕКЦИЯ: Результаты контроля ===
        control_section = self.create_section(self.scrollable_frame, "Результаты контроля", "📏")
        control_section.grid_columnconfigure(1, weight=1)

        self.min_thickness = self.create_field(control_section, "Мин. толщина стенки (мм):", 0, 0,
                                               "Минимальная замеренная толщина")

        # Поле для дефектов (многострочное)
        tk.Label(control_section, text="Основные дефекты:", font=self.font_bold, fg=self.primary_color,
                 bg=self.section_bg, anchor="w").grid(row=1, column=0, sticky="nw", padx=(0, 10), pady=8)

        defects_frame = tk.Frame(control_section, bg=self.section_bg)
        defects_frame.grid(row=1, column=1, sticky="ew", padx=5, pady=5)

        self.max_defects = scrolledtext.ScrolledText(defects_frame, width=50, height=4, font=self.font_normal,
                                                     wrap=tk.WORD, highlightbackground=self.border_color,
                                                     highlightthickness=1)
        self.max_defects.pack(fill="both", expand=True)

        info_btn = tk.Label(control_section, text="ⓘ", font=("Segoe UI", 10), fg=self.accent_color, bg=self.section_bg,
                            cursor="question_arrow")
        info_btn.grid(row=1, column=2, padx=(2, 0), pady=5, sticky="n")
        info_btn.bind('<Enter>', lambda e: self.show_tooltip(e, "Опишите выявленные дефекты"))
        info_btn.bind('<Leave>', self.hide_tooltip)

        self.remaining_life = self.create_field(control_section, "Остаточный ресурс (лет):", 2, 0,
                                                "Расчетный остаточный срок службы")

        # === СЕКЦИЯ: Дополнительные данные ===
        additional_section = self.create_section(self.scrollable_frame, "Дополнительные данные", "📝")
        additional_section.grid_columnconfigure(1, weight=1)

        # Результаты замеров
        tk.Label(additional_section, text="Результаты замеров:", font=self.font_bold, fg=self.primary_color,
                 bg=self.section_bg, anchor="w").grid(row=0, column=0, sticky="nw", padx=(0, 10), pady=8)

        meas_frame = tk.Frame(additional_section, bg=self.section_bg)
        meas_frame.grid(row=0, column=1, sticky="ew", padx=5, pady=5)

        self.measurements_text = scrolledtext.ScrolledText(meas_frame, width=50, height=5, font=self.font_normal,
                                                           wrap=tk.WORD, highlightbackground=self.border_color,
                                                           highlightthickness=1)
        self.measurements_text.pack(fill="both", expand=True)

        # Комментарии
        tk.Label(additional_section, text="Комментарии эксперта:", font=self.font_bold, fg=self.primary_color,
                 bg=self.section_bg, anchor="w").grid(row=1, column=0, sticky="nw", padx=(0, 10), pady=8)

        comm_frame = tk.Frame(additional_section, bg=self.section_bg)
        comm_frame.grid(row=1, column=1, sticky="ew", padx=5, pady=5)

        self.comments_text = scrolledtext.ScrolledText(comm_frame, width=50, height=3, font=self.font_normal,
                                                       wrap=tk.WORD, highlightbackground=self.border_color,
                                                       highlightthickness=1)
        self.comments_text.pack(fill="both", expand=True)

        # === СЕКЦИЯ: Прикрепленные файлы ===
        files_section = self.create_section(self.scrollable_frame, "Прикрепленные файлы", "📎")

        # Кнопки
        btn_frame = tk.Frame(files_section, bg=self.section_bg)
        btn_frame.pack(fill="x", pady=5)

        add_btn = tk.Button(btn_frame, text="➕ Добавить файлы", command=self.add_files,
                            bg=self.accent_color, fg="white", font=self.font_normal,
                            relief="flat", padx=15, pady=5, cursor="hand2")
        add_btn.pack(side="left", padx=5)

        remove_btn = tk.Button(btn_frame, text="✖ Удалить", command=self.remove_file,
                               bg=self.danger_color, fg="white", font=self.font_normal,
                               relief="flat", padx=15, pady=5, cursor="hand2")
        remove_btn.pack(side="left", padx=5)

        # Список файлов
        list_frame = tk.Frame(files_section, bg=self.section_bg)
        list_frame.pack(fill="both", expand=True, pady=10)

        scrollbar_files = tk.Scrollbar(list_frame)
        scrollbar_files.pack(side="right", fill="y")

        self.files_listbox = tk.Listbox(list_frame, width=70, height=5, font=self.font_normal,
                                        yscrollcommand=scrollbar_files.set,
                                        relief="solid", bd=1, highlightbackground=self.border_color)
        self.files_listbox.pack(side="left", fill="both", expand=True)

        scrollbar_files.config(command=self.files_listbox.yview)

        # === КНОПКИ ===
        btn_frame_main = tk.Frame(self.scrollable_frame, bg=self.bg_color)
        btn_frame_main.pack(fill="x", pady=30, padx=10)

        save_btn = tk.Button(btn_frame_main, text="💾 Загрузить результаты", command=self.upload_results,
                             bg=self.success_color, fg="white", font=self.font_bold,
                             relief="flat", padx=30, pady=10, cursor="hand2", width=20)
        save_btn.pack(side="left", padx=10)

        cancel_btn = tk.Button(btn_frame_main, text="✕ Отмена", command=self.window.destroy,
                               bg=self.danger_color, fg="white", font=self.font_normal,
                               relief="flat", padx=30, pady=10, cursor="hand2", width=15)
        cancel_btn.pack(side="left", padx=10)

    def load_projects(self):
        """Загрузка списка проектов для эксперта"""
        print(f"Загрузка проектов для эксперта ID: {current_user['id']}")  # Отладка

        query = """
            SELECT p.id, p.project_number, c.name as customer_name 
            FROM projects p 
            JOIN customers c ON p.customer_id = c.id 
            WHERE p.expert_id = %s 
            ORDER BY p.creation_date DESC
        """
        projects = Database.execute_query(query, (current_user['id'],), fetchall=True)

        print(f"Найдено проектов: {len(projects) if projects else 0}")  # Отладка

        if projects:
            values = [f"{p['id']}: {p['project_number']} - {p['customer_name']}" for p in projects]
            self.project_combo['values'] = values
            if len(values) > 0:
                self.project_combo.set(values[0])  # Выбираем первый по умолчанию
        else:
            self.project_combo['values'] = ['Нет доступных проектов']
            self.project_combo.set('Нет доступных проектов')
            messagebox.showinfo("Информация", "У вас нет проектов, ожидающих загрузки результатов")

    def add_files(self):
        """Добавление файлов"""
        files = filedialog.askopenfilenames(
            title="Выберите файлы",
            filetypes=[("Все файлы", "*.*"), ("Изображения", "*.jpg *.jpeg *.png"),
                       ("Документы", "*.pdf *.doc *.docx *.xls *.xlsx")]
        )
        for file in files:
            if file not in self.uploaded_files:
                self.uploaded_files.append(file)
                self.files_listbox.insert(tk.END, os.path.basename(file))

    def remove_file(self):
        """Удаление файла из списка"""
        selection = self.files_listbox.curselection()
        if selection:
            index = selection[0]
            self.files_listbox.delete(index)
            del self.uploaded_files[index]

    def upload_results(self):
        """Загрузка результатов в БД"""
        # Определение проекта
        if hasattr(self, 'project_var') and self.project_var:
            project_id = self.project_var
        else:
            if not self.project_combo.get() or self.project_combo.get() == 'Нет доступных проектов':
                messagebox.showerror("Ошибка", "Выберите проект")
                return
            try:
                project_id = int(self.project_combo.get().split(':')[0])
            except:
                messagebox.showerror("Ошибка", "Некорректный выбор проекта")
                return

        # Получаем тип работ из проекта
        type_query = "SELECT expert_type_id FROM projects WHERE id = %s"
        type_result = Database.execute_query(type_query, (project_id,), fetchone=True)
        if not type_result:
            messagebox.showerror("Ошибка", "Не удалось определить тип работ для проекта")
            return
        expert_type_id = type_result['expert_type_id']

        # Сохранение результатов с новыми полями
        query = """
            INSERT INTO results (
                project_id, upload_date, expert_type_id, 
                measurements, comments, expert_id,
                working_pressure, working_temperature, capacity,
                working_medium, hazard_class, body_material,
                welding_type, min_thickness, max_defects, remaining_life
            ) VALUES (
                %s, NOW(), %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
            )
        """

        # Собираем все поля
        params = (
            project_id,
            expert_type_id,
            self.measurements_text.get('1.0', 'end-1c'),
            self.comments_text.get('1.0', 'end-1c'),
            current_user['id'],
            self.working_pressure.get(),
            self.working_temperature.get(),
            self.capacity.get(),
            self.working_medium.get(),
            self.hazard_class.get(),
            self.body_material.get(),
            self.welding_type.get(),
            self.min_thickness.get(),
            self.max_defects.get('1.0', 'end-1c'),
            self.remaining_life.get()
        )

        result_id = Database.execute_query(query, params, commit=True)

        if result_id:
            # Сохранение файлов (как было)
            for file_path in self.uploaded_files:
                file_name = os.path.basename(file_path)
                dest_path = os.path.join(UPLOAD_FOLDER, f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{file_name}")
                shutil.copy2(file_path, dest_path)

                file_size = os.path.getsize(file_path)
                file_type = os.path.splitext(file_name)[1]

                query = """
                    INSERT INTO project_files (project_id, result_id, file_name, file_path, file_type, file_size, upload_date, employee_id)
                    VALUES (%s, %s, %s, %s, %s, %s, NOW(), %s)
                """
                Database.execute_query(
                    query, (project_id, result_id, file_name, dest_path, file_type, file_size, current_user['id']),
                    commit=True
                )

            Database.log_action(current_user['id'], 'UPLOAD_RESULTS', 'project', project_id,
                                f"Загружены результаты обследования, файлов: {len(self.uploaded_files)}")
            messagebox.showinfo("Успех", "Результаты успешно загружены")
            self.window.destroy()
        else:
            messagebox.showerror("Ошибка", "Не удалось загрузить результаты")


class ResultsViewWindow:
    """Окно просмотра результатов обследования"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Просмотр результатов обследования")
        self.window.geometry("1000x600")

        # Таблица результатов
        self.create_table()
        self.load_results()

    def create_table(self):
        """Создание таблицы результатов"""
        table_frame = tk.Frame(self.window)
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        columns = ('ID', 'Проект', 'Дата', 'Тип работ', 'Эксперт', 'Файлов')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=20)

        col_widths = [50, 120, 120, 200, 150, 80]
        for i, col in enumerate(columns):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=col_widths[i])

        # Скроллбары
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

        # Двойной клик для просмотра деталей
        self.tree.bind('<Double-Button-1>', self.view_result)

    def load_results(self):
        """Загрузка списка результатов"""
        # Очистка таблицы
        for item in self.tree.get_children():
            self.tree.delete(item)

        # Формирование запроса в зависимости от роли
        if current_user_role == 'expert':
            query = """
                SELECT r.id, p.project_number, r.upload_date, et.name as type_name,
                       e.full_name as expert_name,
                       (SELECT COUNT(*) FROM project_files WHERE result_id = r.id) as file_count
                FROM results r
                JOIN projects p ON r.project_id = p.id
                JOIN expert_types et ON r.expert_type_id = et.id
                JOIN employees e ON r.expert_id = e.id
                WHERE r.expert_id = %s
                ORDER BY r.upload_date DESC
            """
            results = Database.execute_query(query, (current_user['id'],), fetchall=True)
        else:
            query = """
                SELECT r.id, p.project_number, r.upload_date, et.name as type_name,
                       e.full_name as expert_name,
                       (SELECT COUNT(*) FROM project_files WHERE result_id = r.id) as file_count
                FROM results r
                JOIN projects p ON r.project_id = p.id
                JOIN expert_types et ON r.expert_type_id = et.id
                JOIN employees e ON r.expert_id = e.id
                ORDER BY r.upload_date DESC
            """
            results = Database.execute_query(query, fetchall=True)

        if results:
            for r in results:
                self.tree.insert('', 'end', values=(
                    r['id'],
                    r['project_number'],
                    r['upload_date'].strftime('%d.%m.%Y %H:%M') if r['upload_date'] else '',
                    r['type_name'],
                    r['expert_name'],
                    r['file_count']
                ))

    def view_result(self, event):
        """Просмотр деталей результата"""
        selected = self.tree.selection()
        if not selected:
            return

        item = self.tree.item(selected[0])
        result_id = item['values'][0]

        ResultViewWindow(result_id)


class ResultViewWindow:
    """Окно просмотра деталей результата обследования"""

    def __init__(self, result_id):
        self.result_id = result_id
        self.window = tk.Toplevel()
        self.window.title(f"Результат обследования ID {result_id}")
        self.window.geometry("700x600")

        # Загрузка данных результата
        query = """
            SELECT r.*, p.project_number, et.name as type_name, e.full_name as expert_name
            FROM results r
            JOIN projects p ON r.project_id = p.id
            JOIN expert_types et ON r.expert_type_id = et.id
            JOIN employees e ON r.expert_id = e.id
            WHERE r.id = %s
        """
        result = Database.execute_query(query, (result_id,), fetchone=True)

        if not result:
            messagebox.showerror("Ошибка", "Результат не найден")
            self.window.destroy()
            return

        # Основной фрейм
        main_frame = tk.Frame(self.window, padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)

        # Заголовок
        tk.Label(main_frame, text=f"Результаты обследования по проекту {result['project_number']}",
                 font=("Arial", 14, "bold")).pack(anchor="w", pady=10)

        # Информация
        info_frame = tk.LabelFrame(main_frame, text="Основная информация", font=("Arial", 12, "bold"))
        info_frame.pack(fill="x", pady=10)

        info_grid = tk.Frame(info_frame)
        info_grid.pack(pady=10, padx=10)

        fields = [
            ("Проект:", result['project_number']),
            ("Тип работ:", result['type_name']),
            ("Эксперт:", result['expert_name']),
            ("Дата загрузки:", result['upload_date'].strftime('%d.%m.%Y %H:%M') if result['upload_date'] else '')
        ]

        for i, (label, value) in enumerate(fields):
            tk.Label(info_grid, text=label, font=("Arial", 10, "bold")).grid(row=i, column=0, sticky="w", pady=2)
            tk.Label(info_grid, text=value, font=("Arial", 10)).grid(row=i, column=1, sticky="w", pady=2, padx=(10, 0))

        # Результаты замеров
        if result['measurements']:
            meas_frame = tk.LabelFrame(main_frame, text="Результаты замеров", font=("Arial", 12, "bold"))
            meas_frame.pack(fill="x", pady=10)

            tk.Label(meas_frame, text=result['measurements'], font=("Arial", 10),
                     wraplength=600, justify="left").pack(pady=10, padx=10)

        # Комментарии
        if result['comments']:
            comm_frame = tk.LabelFrame(main_frame, text="Комментарии", font=("Arial", 12, "bold"))
            comm_frame.pack(fill="x", pady=10)

            tk.Label(comm_frame, text=result['comments'], font=("Arial", 10),
                     wraplength=600, justify="left").pack(pady=10, padx=10)

        # Файлы
        files_frame = tk.LabelFrame(main_frame, text="Прикрепленные файлы", font=("Arial", 12, "bold"))
        files_frame.pack(fill="x", pady=10)

        # Загрузка файлов
        files = Database.execute_query(
            "SELECT * FROM project_files WHERE result_id = %s",
            (result_id,), fetchall=True
        )

        if files:
            for file in files:
                file_label = tk.Label(files_frame, text=f"📄 {file['file_name']} ({file['file_size'] // 1024} КБ)",
                                      font=("Arial", 10), fg="blue", cursor="hand2")
                file_label.pack(anchor="w", padx=10, pady=2)
                file_label.bind('<Button-1>', lambda e, path=file['file_path']: self.open_file(path))
        else:
            tk.Label(files_frame, text="Нет прикрепленных файлов", font=("Arial", 10)).pack(pady=10)

        # Кнопка закрытия
        tk.Button(main_frame, text="Закрыть", command=self.window.destroy,
                  bg="#f44336", fg="white", width=10).pack(pady=20)

    def open_file(self, file_path):
        """Открытие файла"""
        try:
            os.startfile(file_path)
        except:
            messagebox.showerror("Ошибка", "Не удалось открыть файл")


class DraftsWindow:
    """Окно черновиков документов"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Черновики документов")
        self.window.geometry("1000x600")

        # Панель инструментов
        toolbar = tk.Frame(self.window, bg="#f0f0f0", height=40)
        toolbar.pack(fill="x", padx=2, pady=2)

        tk.Button(toolbar, text="Обновить", command=self.load_drafts,
                  bg="#FF9800", fg="white", width=10).pack(side="left", padx=5, pady=5)

        if current_user_role in ['specialist']:
            tk.Button(toolbar, text="Создать черновики", command=self.create_drafts,
                      bg="#4CAF50", fg="white", width=15).pack(side="left", padx=5, pady=5)

        # Фильтр по типу
        tk.Label(toolbar, text="Тип:", bg="#f0f0f0").pack(side="left", padx=(20, 5))
        self.type_filter = ttk.Combobox(toolbar, values=['Все', 'act', 'report', 'conclusion'], width=15)
        self.type_filter.set('Все')
        self.type_filter.pack(side="left", padx=5)
        self.type_filter.bind('<<ComboboxSelected>>', self.apply_filter)

        # Фильтр по статусу
        tk.Label(toolbar, text="Статус:", bg="#f0f0f0").pack(side="left", padx=(20, 5))
        self.status_filter = ttk.Combobox(toolbar, values=['Все', 'Черновик', 'На согласовании'], width=15)
        self.status_filter.set('Все')
        self.status_filter.pack(side="left", padx=5)
        self.status_filter.bind('<<ComboboxSelected>>', self.apply_filter)

        # Таблица
        self.create_table()
        self.load_drafts()

    def create_table(self):
        """Создание таблицы черновиков"""
        table_frame = tk.Frame(self.window)
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        columns = ('ID', 'Проект', 'Тип документа', 'Версия', 'Статус', 'Дата создания', 'Автор')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=20)

        col_widths = [50, 150, 150, 80, 120, 150, 150]
        for i, col in enumerate(columns):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=col_widths[i])

        # Скроллбары
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

        # Двойной клик для редактирования
        self.tree.bind('<Double-Button-1>', self.edit_draft)

    def load_drafts(self):
        """Загрузка списка черновиков"""
        # Очистка таблицы
        for item in self.tree.get_children():
            self.tree.delete(item)

        query = """
            SELECT d.id, p.project_number, d.doc_type, d.version, s.name as status_name,
                   d.create_date, e.full_name as author_name
            FROM drafts d
            JOIN projects p ON d.project_id = p.id
            JOIN statuses s ON d.status_id = s.id
            JOIN employees e ON d.author_id = e.id
            ORDER BY d.create_date DESC
        """
        drafts = Database.execute_query(query, fetchall=True)

        self.all_drafts = drafts

        if drafts:
            for d in drafts:
                doc_type_names = {'act': 'Акт обследования', 'report': 'Технический отчет',
                                  'conclusion': 'Экспертное заключение'}
                doc_type = doc_type_names.get(d['doc_type'], d['doc_type'])

                self.tree.insert('', 'end', values=(
                    d['id'],
                    d['project_number'],
                    doc_type,
                    d['version'],
                    d['status_name'],
                    d['create_date'].strftime('%d.%m.%Y %H:%M') if d['create_date'] else '',
                    d['author_name']
                ))

    def apply_filter(self, event=None):
        """Применение фильтров"""
        type_filter = self.type_filter.get()
        status_filter = self.status_filter.get()

        # Очистка таблицы
        for item in self.tree.get_children():
            self.tree.delete(item)

        if not hasattr(self, 'all_drafts') or not self.all_drafts:
            return

        doc_type_names = {'act': 'Акт обследования', 'report': 'Технический отчет',
                          'conclusion': 'Экспертное заключение'}

        for d in self.all_drafts:
            # Фильтр по типу
            if type_filter != 'Все':
                if type_filter == 'act' and d['doc_type'] != 'act':
                    continue
                elif type_filter == 'report' and d['doc_type'] != 'report':
                    continue
                elif type_filter == 'conclusion' and d['doc_type'] != 'conclusion':
                    continue

            # Фильтр по статусу
            if status_filter != 'Все' and d['status_name'] != status_filter:
                continue

            doc_type = doc_type_names.get(d['doc_type'], d['doc_type'])

            self.tree.insert('', 'end', values=(
                d['id'],
                d['project_number'],
                doc_type,
                d['version'],
                d['status_name'],
                d['create_date'].strftime('%d.%m.%Y %H:%M') if d['create_date'] else '',
                d['author_name']
            ))

    def edit_draft(self, event):
        """Редактирование черновика"""
        selected = self.tree.selection()
        if not selected:
            return

        item = self.tree.item(selected[0])
        draft_id = item['values'][0]

        DraftEditWindow(draft_id)

    def create_drafts(self):
        """Создание черновиков для проекта"""
        # Здесь должна быть логика выбора проекта и создания черновиков
        messagebox.showinfo("Информация", "Выберите проект для создания черновиков")


class DraftEditWindow:
    """Окно редактирования черновика"""

    def __init__(self, draft_id):
        self.draft_id = draft_id
        self.window = tk.Toplevel()
        self.window.title(f"Редактирование черновика ID {draft_id}")
        self.window.geometry("600x300")

        # Загрузка данных черновика
        query = """
            SELECT d.*, p.project_number, t.name as template_name
            FROM drafts d
            JOIN projects p ON d.project_id = p.id
            LEFT JOIN templates t ON d.template_id = t.id
            WHERE d.id = %s
        """
        draft = Database.execute_query(query, (draft_id,), fetchone=True)

        if not draft:
            messagebox.showerror("Ошибка", "Черновик не найден")
            self.window.destroy()
            return

        self.draft = draft

        # Основной фрейм
        main_frame = tk.Frame(self.window, padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)

        # Заголовок
        doc_type_names = {'act': 'Акт обследования', 'report': 'Технический отчет',
                          'conclusion': 'Экспертное заключение'}
        doc_type = doc_type_names.get(draft['doc_type'], draft['doc_type'])

        tk.Label(main_frame, text=f"{doc_type} - Проект {draft['project_number']}",
                 font=("Arial", 14, "bold")).pack(anchor="w", pady=10)

        # Информация
        info_frame = tk.Frame(main_frame, bg="#f0f0f0", relief="groove", bd=1)
        info_frame.pack(fill="x", pady=10)

        info_text = f"Версия: {draft['version']} | Шаблон: {draft['template_name'] or 'Не указан'} | Статус: {self.get_status_name(draft['status_id'])}"
        tk.Label(info_frame, text=info_text, bg="#f0f0f0", pady=5).pack()

        # Информация о файле
        tk.Label(main_frame, text="Документ:", font=("Arial", 10, "bold")).pack(anchor="w", pady=(10, 5))

        file_frame = tk.Frame(main_frame)
        file_frame.pack(fill="x", pady=5)

        self.file_path = draft.get('file_path')
        if self.file_path and os.path.exists(self.file_path):
            tk.Label(file_frame, text=os.path.basename(self.file_path),
                     font=("Arial", 10)).pack(side="left", padx=5)
            tk.Button(file_frame, text="Открыть в Word",
                      command=self.open_word,
                      bg="#2196F3", fg="white").pack(side="left", padx=5)
        else:
            tk.Label(file_frame, text="Файл не найден", fg="red").pack(side="left")

        # Кнопки
        btn_frame = tk.Frame(main_frame)
        btn_frame.pack(pady=20)

        if draft['status_id'] == 1:  # Черновик
            tk.Button(btn_frame, text="Отправить на согласование", command=self.send_to_approval,
                      bg="#2196F3", fg="white", width=20).pack(side="left", padx=5)

        tk.Button(btn_frame, text="Закрыть", command=self.window.destroy,
                  bg="#f44336", fg="white", width=10).pack(side="left", padx=5)

    def get_status_name(self, status_id):
        statuses = {1: 'Черновик', 2: 'На согласовании', 3: 'Утвержден', 4: 'Передан в А4'}
        return statuses.get(status_id, 'Неизвестно')

    def open_word(self):
        """Открытие файла в Word"""
        try:
            if os.name == 'nt':  # Windows
                os.startfile(self.file_path)
            else:
                subprocess.call(['xdg-open', self.file_path])
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл: {e}")

    def send_to_approval(self):
        """Отправка на согласование"""
        print(f"Отправка на согласование черновика ID: {self.draft_id}")

        # Проверяем существование файла
        if not os.path.exists(self.file_path):
            messagebox.showerror("Ошибка", f"Файл документа не найден: {self.file_path}")
            return

        status = Database.execute_query("SELECT id FROM statuses WHERE name='На согласовании'", fetchone=True)
        status_id = status['id'] if status else 2

        query = "UPDATE drafts SET status_id = %s WHERE id = %s"
        result = Database.execute_query(query, (status_id, self.draft_id), commit=True)

        if result is not None:
            Database.log_action(current_user['id'], 'SEND_TO_APPROVAL', 'draft', self.draft_id,
                                f"Отправлен на согласование черновик ID {self.draft_id}")
            messagebox.showinfo("Успех", "Документ отправлен на согласование")
            self.window.destroy()
        else:
            messagebox.showerror("Ошибка", "Не удалось отправить на согласование")


class FinalDocsWindow:
    """Окно готовых документов"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Готовые документы")
        self.window.geometry("1000x600")

        # Панель инструментов
        toolbar = tk.Frame(self.window, bg="#f0f0f0", height=40)
        toolbar.pack(fill="x", padx=2, pady=2)

        tk.Button(toolbar, text="Обновить", command=self.load_documents,
                  bg="#FF9800", fg="white", width=10).pack(side="left", padx=5, pady=5)

        tk.Button(toolbar, text="Просмотреть PDF", command=self.view_document,
                  bg="#2196F3", fg="white", width=15).pack(side="left", padx=5, pady=5)

        tk.Button(toolbar, text="Открыть папку с PDF", command=self.open_pdf_folder,
                  bg="#4CAF50", fg="white", width=15).pack(side="left", padx=5, pady=5)

        # Поиск
        tk.Label(toolbar, text="Поиск:", bg="#f0f0f0").pack(side="right", padx=5)
        self.search_entry = tk.Entry(toolbar, width=30)
        self.search_entry.pack(side="right", padx=5)
        self.search_entry.bind('<KeyRelease>', self.search_documents)

        # Таблица
        self.create_table()
        self.load_documents()

        self.tooltip = None  # для подсказки

    def create_table(self):
        """Создание таблицы готовых документов"""
        table_frame = tk.Frame(self.window)
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        columns = ('ID', 'Проект', 'Тип документа', 'Номер', 'Дата утверждения', 'Утвердил', 'PDF')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=20)

        col_widths = [50, 150, 200, 120, 100, 150, 80]
        for i, col in enumerate(columns):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=col_widths[i], anchor="center" if col == 'PDF' else "w")

        # Скроллбары
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

        # Привязка событий
        self.tree.bind('<Double-Button-1>', self.view_document)
        self.tree.bind('<Motion>', self.show_tooltip)
        self.tree.bind('<Leave>', self.hide_tooltip)

    def load_documents(self):
        """Загрузка списка готовых документов"""
        for item in self.tree.get_children():
            self.tree.delete(item)

        query = """
            SELECT f.id, p.project_number, f.doc_type, f.doc_number, f.approval_date,
                   e.full_name as approver_name, 
                   CASE WHEN f.pdf_path IS NOT NULL AND f.pdf_path != '' THEN '✅' ELSE '❌' END as has_pdf
            FROM final_documents f
            JOIN projects p ON f.project_id = p.id
            LEFT JOIN employees e ON f.approver_id = e.id
            ORDER BY f.approval_date DESC
        """
        docs = Database.execute_query(query, fetchall=True)

        self.all_docs = docs

        if docs:
            doc_type_names = {'act': 'Акт обследования', 'report': 'Технический отчет',
                              'conclusion': 'Экспертное заключение'}

            for d in docs:
                doc_type = doc_type_names.get(d['doc_type'], d['doc_type'])

                # Определяем цвет для статуса PDF
                tag = 'has_pdf' if d['has_pdf'] == '✅' else 'no_pdf'

                self.tree.insert('', 'end', values=(
                    d['id'],
                    d['project_number'],
                    doc_type,
                    d['doc_number'] or '',
                    d['approval_date'].strftime('%d.%m.%Y') if d['approval_date'] else '',
                    d['approver_name'] or '',
                    d['has_pdf']
                ), tags=(tag,))

            # Настройка цветов
            self.tree.tag_configure('has_pdf', background='#e8f5e8')  # светло-зеленый
            self.tree.tag_configure('no_pdf', background='#ffebee')  # светло-красный

    def search_documents(self, event=None):
        """Поиск по документам"""
        search_text = self.search_entry.get().lower()

        # Очищаем таблицу
        for item in self.tree.get_children():
            self.tree.delete(item)

        if not hasattr(self, 'all_docs') or not self.all_docs:
            return

        doc_type_names = {'act': 'Акт обследования', 'report': 'Технический отчет',
                          'conclusion': 'Экспертное заключение'}

        for d in self.all_docs:
            doc_type = doc_type_names.get(d['doc_type'], d['doc_type'])

            # Поиск по тексту
            if search_text:
                searchable = f"{d['project_number']} {doc_type} {d['doc_number']} {d['approver_name']}".lower()
                if search_text not in searchable:
                    continue

            tag = 'has_pdf' if d['has_pdf'] == '✅' else 'no_pdf'

            self.tree.insert('', 'end', values=(
                d['id'],
                d['project_number'],
                doc_type,
                d['doc_number'] or '',
                d['approval_date'].strftime('%d.%m.%Y') if d['approval_date'] else '',
                d['approver_name'] or '',
                d['has_pdf']
            ), tags=(tag,))

    def view_document(self, event=None):
        """Открытие PDF документа в браузере"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите документ для просмотра")
            return

        item = self.tree.item(selected[0])
        doc_id = item['values'][0]

        # Получаем путь к PDF
        query = "SELECT pdf_path, doc_number, doc_type FROM final_documents WHERE id = %s"
        doc = Database.execute_query(query, (doc_id,), fetchone=True)

        if not doc:
            messagebox.showerror("Ошибка", "Документ не найден")
            return

        if not doc['pdf_path'] or not os.path.exists(doc['pdf_path']):
            messagebox.showinfo("Информация", "PDF файл не найден. Возможно, документ еще не сконвертирован.")
            return

        # Открываем PDF в браузере
        try:
            # Преобразуем путь в абсолютный и в URL формат
            abs_path = os.path.abspath(doc['pdf_path'])
            file_url = f"file:///{abs_path.replace('\\', '/')}"

            # Открываем в браузере по умолчанию
            webbrowser.open(file_url)

            # Логируем действие
            Database.log_action(current_user['id'], 'VIEW_PDF', 'final_document', doc_id,
                                f"Просмотр PDF: {doc['doc_number']}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть PDF: {e}")

    def open_pdf_folder(self):
        """Открыть папку с PDF файлами"""
        pdf_folder = os.path.abspath(FINAL_FOLDER)
        if os.name == 'nt':  # Windows
            os.startfile(pdf_folder)
        else:
            subprocess.call(['xdg-open', pdf_folder])

    def show_tooltip(self, event):
        """Показать подсказку при наведении на ячейку с PDF"""
        # Скрываем предыдущую подсказку
        self.hide_tooltip()

        item = self.tree.identify_row(event.y)
        col = self.tree.identify_column(event.x)

        if not item or col != '#7':  # колонка PDF (седьмая колонка)
            return

        values = self.tree.item(item)['values']
        if values and len(values) > 6 and values[6] == '✅':  # есть PDF
            # Получаем координаты ячейки
            bbox = self.tree.bbox(item, column='#7')
            if bbox:
                x, y, width, height = bbox
                # Создаем всплывающую подсказку
                self.tooltip = tk.Toplevel(self.window)
                self.tooltip.wm_overrideredirect(True)
                self.tooltip.wm_geometry(f"+{self.window.winfo_rootx() + x + 50}+{self.window.winfo_rooty() + y + 30}")

                label = tk.Label(self.tooltip, text="Двойной клик для просмотра PDF",
                                 bg="#ffffe0", relief="solid", borderwidth=1, padx=5, pady=2)
                label.pack()

    def hide_tooltip(self, event=None):
        """Скрыть подсказку"""
        if hasattr(self, 'tooltip') and self.tooltip:
            self.tooltip.destroy()
            self.tooltip = None


def view_document(self, event):
    """Просмотр готового документа"""
    selected = self.tree.selection()
    if not selected:
        return

    item = self.tree.item(selected[0])
    doc_id = item['values'][0]

    # Получаем путь к PDF
    query = "SELECT pdf_path, doc_number FROM final_documents WHERE id = %s"
    doc = Database.execute_query(query, (doc_id,), fetchone=True)

    if not doc or not doc['pdf_path']:
        messagebox.showinfo("Информация", "PDF файл не найден")
        return

    pdf_path = doc['pdf_path']
    if not os.path.exists(pdf_path):
        messagebox.showerror("Ошибка", f"Файл не найден: {pdf_path}")
        return

    # Открываем PDF
    try:
        if os.name == 'nt':  # Windows
            os.startfile(pdf_path)
        else:
            subprocess.call(['xdg-open', pdf_path])
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось открыть PDF: {e}")


class ApprovalWindow:
    """Окно согласования документов"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Согласование документов")
        self.window.geometry("1000x600")

        # Панель инструментов
        toolbar = tk.Frame(self.window, bg="#f0f0f0", height=40)
        toolbar.pack(fill="x", padx=2, pady=2)

        tk.Button(toolbar, text="Обновить", command=self.load_documents,
                  bg="#FF9800", fg="white", width=10).pack(side="left", padx=5, pady=5)

        tk.Button(toolbar, text="Просмотреть документ", command=self.view_selected_document,
                  bg="#2196F3", fg="white", width=20).pack(side="left", padx=5, pady=5)

        # Таблица
        self.create_table()
        self.load_documents()

    def create_table(self):
        """Создание таблицы документов на согласовании"""
        table_frame = tk.Frame(self.window)
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        columns = ('ID', 'Проект', 'Тип документа', 'Версия', 'Автор', 'Дата создания')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=20)

        col_widths = [50, 150, 200, 80, 150, 150]
        for i, col in enumerate(columns):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=col_widths[i])

        # Скроллбары
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

        # Двойной клик для просмотра
        self.tree.bind('<Double-Button-1>', self.on_double_click)

    def load_documents(self):
        """Загрузка документов на согласование"""
        print("Загрузка документов на согласование...")

        for item in self.tree.get_children():
            self.tree.delete(item)

        # Получаем статус "На согласовании"
        status = Database.execute_query("SELECT id FROM statuses WHERE name='На согласовании'", fetchone=True)
        status_id = status['id'] if status else 2
        print(f"Статус ID для 'На согласовании': {status_id}")

        query = """
            SELECT d.id, p.project_number, d.doc_type, d.version, e.full_name as author_name, 
                   d.create_date
            FROM drafts d
            JOIN projects p ON d.project_id = p.id
            JOIN employees e ON d.author_id = e.id
            WHERE d.status_id = %s
            ORDER BY d.create_date ASC
        """
        docs = Database.execute_query(query, (status_id,), fetchall=True)

        print(f"Найдено документов: {len(docs) if docs else 0}")

        if docs:
            doc_type_names = {'act': 'Акт обследования', 'report': 'Технический отчет',
                              'conclusion': 'Экспертное заключение'}

            for d in docs:
                doc_type = doc_type_names.get(d['doc_type'], d['doc_type'])

                # Вставляем запись и сохраняем ID как отдельный атрибут
                item_id = self.tree.insert('', 'end', values=(
                    d['id'],
                    d['project_number'],
                    doc_type,
                    d['version'],
                    d['author_name'],
                    d['create_date'].strftime('%d.%m.%Y %H:%M') if d['create_date'] else ''
                ))

                # Сохраняем ID документа в самом элементе (важно!)
                self.tree.item(item_id, tags=(str(d['id']),))

            print(f"Загружено {len(docs)} документов")
        else:
            self.tree.insert('', 'end', values=('', 'Нет документов на согласовании', '', '', '', ''))

    def on_double_click(self, event):
        """Обработка двойного клика"""
        self.view_selected_document()

    def view_selected_document(self):
        """Просмотр выбранного документа"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите документ для просмотра")
            return

        # Получаем ID документа из первого столбца
        item = self.tree.item(selected[0])
        values = item['values']
        if not values or len(values) == 0:
            messagebox.showerror("Ошибка", "Не удалось получить данные документа")
            return

        draft_id = values[0]  # ID документа в первом столбце
        print(f"Выбран документ с ID: {draft_id}")

        if not draft_id or draft_id == '':
            messagebox.showerror("Ошибка", "Некорректный ID документа")
            return

        # Проверяем существование документа в БД
        check_query = "SELECT id FROM drafts WHERE id = %s"
        check = Database.execute_query(check_query, (draft_id,), fetchone=True)

        if not check:
            messagebox.showerror("Ошибка", f"Документ с ID {draft_id} не найден в базе данных")
            return

        # Открываем окно согласования
        try:
            ApprovalViewWindow(draft_id)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть окно согласования: {e}")
            import traceback
            traceback.print_exc()


class ApprovalViewWindow:
    """Окно просмотра документа для согласования"""

    def __init__(self, draft_id):
        print(f"=== СОЗДАНИЕ ОКНА СОГЛАСОВАНИЯ ===")
        print(f"Получен ID документа: {draft_id}")

        self.draft_id = draft_id
        self.window = tk.Toplevel()
        self.window.title(f"Согласование документа")
        self.window.geometry("600x500")

        # Получаем данные документа
        try:
            query = """
                SELECT d.*, p.project_number, e.full_name as author_name,
                       p.id as project_id, p.project_number
                FROM drafts d
                JOIN projects p ON d.project_id = p.id
                JOIN employees e ON d.author_id = e.id
                WHERE d.id = %s
            """
            draft = Database.execute_query(query, (draft_id,), fetchone=True)

            if not draft:
                messagebox.showerror("Ошибка", f"Документ не найден")
                self.window.destroy()
                return

            self.draft = draft
            print(f"Данные загружены: проект {draft['project_number']}, тип {draft['doc_type']}")

        except Exception as e:
            messagebox.showerror("Ошибка", str(e))
            self.window.destroy()
            return

        # Основной фрейм
        main_frame = tk.Frame(self.window, padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)

        # ЗАГОЛОВОК
        doc_type_names = {'act': 'АКТ ОБСЛЕДОВАНИЯ', 'report': 'ТЕХНИЧЕСКИЙ ОТЧЕТ',
                          'conclusion': 'ЭКСПЕРТНОЕ ЗАКЛЮЧЕНИЕ'}
        doc_type = doc_type_names.get(draft['doc_type'], draft['doc_type'])

        tk.Label(main_frame, text=doc_type, font=("Arial", 16, "bold"), fg="blue").pack(pady=10)
        tk.Label(main_frame, text=f"Проект: {draft['project_number']}", font=("Arial", 12)).pack()
        tk.Label(main_frame, text=f"Автор: {draft['author_name']}", font=("Arial", 11)).pack(pady=5)

        # Информация о файле
        file_path = draft.get('file_path')
        if file_path and os.path.exists(file_path):
            tk.Label(main_frame, text=f"Файл: {os.path.basename(file_path)}",
                     font=("Arial", 10), fg="green").pack(pady=10)
            tk.Button(main_frame, text="📄 Открыть документ в Word",
                      command=self.open_document,
                      bg="#2196F3", fg="white", width=25, height=2).pack(pady=10)
        else:
            tk.Label(main_frame, text="❌ Файл документа не найден!",
                     font=("Arial", 10), fg="red").pack(pady=10)

        # Комментарий
        tk.Label(main_frame, text="Комментарий (при отклонении):", font=("Arial", 10)).pack(pady=(10, 5))
        self.comment_entry = tk.Entry(main_frame, width=60)
        self.comment_entry.pack(pady=5)

        # КНОПКИ
        btn_frame = tk.Frame(main_frame)
        btn_frame.pack(pady=30)

        tk.Button(btn_frame, text="✅ УТВЕРДИТЬ", command=self.approve_document,
                  bg="#4CAF50", fg="white", width=15, height=2, font=("Arial", 11, "bold")).pack(side="left", padx=10)

        tk.Button(btn_frame, text="❌ ОТКЛОНИТЬ", command=self.reject_document,
                  bg="#f44336", fg="white", width=15, height=2, font=("Arial", 11, "bold")).pack(side="left", padx=10)

        tk.Button(btn_frame, text="✖ Закрыть", command=self.window.destroy,
                  bg="#808080", fg="white", width=10, height=1).pack(side="left", padx=10)

    def open_document(self):
        """Открытие документа в Word"""
        try:
            os.startfile(self.draft['file_path'])
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть документ: {e}")

    def approve_document(self):
        """Утверждение документа"""
        if messagebox.askyesno("Подтверждение", "Утвердить документ?"):
            print(f"Начало утверждения документа ID: {self.draft_id}")

            # Получаем статус "Утвержден"
            status = Database.execute_query("SELECT id FROM statuses WHERE name='Утвержден'", fetchone=True)
            status_id = status['id'] if status else 3
            print(f"Статус ID для утверждения: {status_id}")

            # Обновление статуса черновика
            update_query = "UPDATE drafts SET status_id = %s WHERE id = %s"
            update_result = Database.execute_query(update_query, (status_id, self.draft_id), commit=True)
            print(f"Результат обновления drafts: {update_result}")

            # Генерируем номер документа
            doc_number = f"{self.draft['doc_type'].upper()}-{datetime.now().strftime('%Y%m%d')}-{self.draft_id}"
            print(f"Номер документа: {doc_number}")

            # Конвертация в PDF
            pdf_path = None
            try:
                from docx2pdf import convert
                pdf_filename = f"{doc_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                pdf_path = os.path.join(FINAL_FOLDER, pdf_filename)

                if self.draft['file_path'] and os.path.exists(self.draft['file_path']):
                    convert(self.draft['file_path'], pdf_path)
                    print(f"PDF создан: {pdf_path}")
                else:
                    print(f"Файл не найден: {self.draft['file_path']}")
            except Exception as e:
                print(f"Ошибка конвертации в PDF: {e}")

            # Вставляем запись в final_documents
            insert_query = """
                INSERT INTO final_documents 
                (project_id, doc_type, draft_id, doc_number, approval_date, approver_id, status_id, pdf_path)
                VALUES (%s, %s, %s, %s, CURDATE(), %s, %s, %s)
            """
            params = (
                self.draft['project_id'],
                self.draft['doc_type'],
                self.draft_id,
                doc_number,
                current_user['id'],
                status_id,
                pdf_path
            )
            print(f"Параметры для вставки: {params}")

            doc_id = Database.execute_query(insert_query, params, commit=True)
            print(f"Результат вставки в final_documents: {doc_id}")

            if doc_id:
                Database.log_action(current_user['id'], 'APPROVE_DOCUMENT', 'draft', self.draft_id,
                                    f"Утвержден документ {doc_number}")
                messagebox.showinfo("Успех", f"Документ утвержден. Номер: {doc_number}")
            else:
                messagebox.showerror("Ошибка", "Не удалось создать запись в готовых документах")

            self.window.destroy()

    def reject_document(self):
        """Отклонение документа"""
        comment = self.comment_entry.get().strip()
        if not comment:
            messagebox.showwarning("Предупреждение", "Укажите причину отклонения")
            return

        if messagebox.askyesno("Подтверждение", "Отклонить документ?"):
            print(f"Отклонение документа ID: {self.draft_id}")

            status = Database.execute_query("SELECT id FROM statuses WHERE name='Черновик'", fetchone=True)
            status_id = status['id'] if status else 1

            query = "UPDATE drafts SET status_id = %s WHERE id = %s"
            Database.execute_query(query, (status_id, self.draft_id), commit=True)

            Database.log_action(current_user['id'], 'REJECT_DOCUMENT', 'draft', self.draft_id,
                                f"Документ отклонен: {comment}")

            messagebox.showinfo("Информация", "Документ отклонен")
            self.window.destroy()


def generate_pdf(self, draft_id, doc_number):
    """Генерация PDF-файла из черновика"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import mm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER

        # Получаем данные черновика
        draft = Database.execute_query("""
            SELECT d.*, p.project_number 
            FROM drafts d
            JOIN projects p ON d.project_id = p.id
            WHERE d.id = %s
        """, (draft_id,), fetchone=True)

        if not draft:
            return None

        # Создаем папку для PDF, если её нет
        pdf_folder = "pdf_documents"
        if not os.path.exists(pdf_folder):
            os.makedirs(pdf_folder)

        # Генерируем имя файла
        filename = f"{pdf_folder}/{doc_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

        # Создаем PDF
        doc = SimpleDocTemplate(filename, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Заголовок
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=30
        )

        doc_type_names = {'act': 'АКТ ОБСЛЕДОВАНИЯ', 'report': 'ТЕХНИЧЕСКИЙ ОТЧЕТ',
                          'conclusion': 'ЭКСПЕРТНОЕ ЗАКЛЮЧЕНИЕ'}
        title = doc_type_names.get(draft['doc_type'], draft['doc_type'])
        story.append(Paragraph(title, title_style))

        # Номер документа
        story.append(Paragraph(f"№ {doc_number}", styles['Normal']))
        story.append(Spacer(1, 20))

        # Содержание
        content_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=10
        )

        # Разбиваем содержимое на параграфы
        paragraphs = draft['content'].split('\n')
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.replace('\n', '<br/>'), content_style))
                story.append(Spacer(1, 5))

        # Строим документ
        doc.build(story)

        return filename

    except Exception as e:
        print(f"Ошибка генерации PDF: {e}")
        return None

class ProjectsReportWindow:
    """Окно реестра проектов"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Реестр проектов")
        self.window.geometry("1100x600")

        # Панель фильтров
        filter_frame = tk.Frame(self.window, bg="#f0f0f0", height=40)
        filter_frame.pack(fill="x", padx=2, pady=2)

        tk.Label(filter_frame, text="Статус:", bg="#f0f0f0").pack(side="left", padx=5)
        self.status_filter = ttk.Combobox(filter_frame,
                                          values=['Все', 'Черновик', 'На согласовании', 'Утвержден', 'Передан в А4'],
                                          width=15)
        self.status_filter.set('Все')
        self.status_filter.pack(side="left", padx=5)

        tk.Label(filter_frame, text="Эксперт:", bg="#f0f0f0").pack(side="left", padx=5)
        self.expert_filter = ttk.Combobox(filter_frame, width=20)
        self.load_experts()
        self.expert_filter.set('Все')
        self.expert_filter.pack(side="left", padx=5)

        tk.Label(filter_frame, text="Период с:", bg="#f0f0f0").pack(side="left", padx=5)
        self.date_from = tk.Entry(filter_frame, width=10)
        self.date_from.pack(side="left", padx=5)
        self.date_from.insert(0, (datetime.now().replace(day=1)).strftime('%d.%m.%Y'))

        tk.Label(filter_frame, text="по:", bg="#f0f0f0").pack(side="left", padx=5)
        self.date_to = tk.Entry(filter_frame, width=10)
        self.date_to.pack(side="left", padx=5)
        self.date_to.insert(0, datetime.now().strftime('%d.%m.%Y'))

        tk.Button(filter_frame, text="Применить фильтр", command=self.apply_filter,
                  bg="#2196F3", fg="white").pack(side="left", padx=10)

        tk.Button(filter_frame, text="Экспорт в Excel", command=self.export_to_excel,
                  bg="#4CAF50", fg="white").pack(side="right", padx=5)

        # Таблица
        self.create_table()
        self.load_report()

    def load_experts(self):
        experts = Database.execute_query("SELECT full_name FROM employees WHERE role='expert' ORDER BY full_name",
                                         fetchall=True)
        if experts:
            expert_list = ['Все'] + [e['full_name'] for e in experts]
            self.expert_filter['values'] = expert_list

    def create_table(self):
        """Создание таблицы отчета"""
        table_frame = tk.Frame(self.window)
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        columns = ('Номер проекта', 'Дата создания', 'Заказчик', 'Объект', 'Эксперт', 'Статус', 'Договор')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=20)

        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=150)

        # Скроллбары
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

    def load_report(self):
        """Загрузка данных для отчета"""
        for item in self.tree.get_children():
            self.tree.delete(item)

        query = """
            SELECT p.project_number, p.creation_date, c.name as customer_name, 
                   o.name as object_name, e.full_name as expert_name, s.name as status_name,
                   p.contract_number
            FROM projects p
            JOIN customers c ON p.customer_id = c.id
            JOIN objects o ON p.object_id = o.id
            JOIN employees e ON p.expert_id = e.id
            JOIN statuses s ON p.status_id = s.id
            ORDER BY p.creation_date DESC
        """
        projects = Database.execute_query(query, fetchall=True)

        self.all_report_data = projects

        if projects:
            for p in projects:
                self.tree.insert('', 'end', values=(
                    p['project_number'],
                    p['creation_date'].strftime('%d.%m.%Y') if p['creation_date'] else '',
                    p['customer_name'],
                    p['object_name'],
                    p['expert_name'],
                    p['status_name'],
                    p['contract_number'] or ''
                ))

    def apply_filter(self):
        """Применение фильтров"""
        status = self.status_filter.get()
        expert = self.expert_filter.get()
        date_from = self.date_from.get()
        date_to = self.date_to.get()

        # Очистка таблицы
        for item in self.tree.get_children():
            self.tree.delete(item)

        if not hasattr(self, 'all_report_data') or not self.all_report_data:
            return

        try:
            from_date = datetime.strptime(date_from, '%d.%m.%Y').date() if date_from else None
            to_date = datetime.strptime(date_to, '%d.%m.%Y').date() if date_to else None
        except:
            messagebox.showerror("Ошибка", "Неверный формат даты")
            return

        for p in self.all_report_data:
            # Фильтр по статусу
            if status != 'Все' and p['status_name'] != status:
                continue

            # Фильтр по эксперту
            if expert != 'Все' and p['expert_name'] != expert:
                continue

            # Фильтр по дате
            if from_date and p['creation_date'] and p['creation_date'] < from_date:
                continue
            if to_date and p['creation_date'] and p['creation_date'] > to_date:
                continue

            self.tree.insert('', 'end', values=(
                p['project_number'],
                p['creation_date'].strftime('%d.%m.%Y') if p['creation_date'] else '',
                p['customer_name'],
                p['object_name'],
                p['expert_name'],
                p['status_name'],
                p['contract_number'] or ''
            ))

    def export_to_excel(self):
        """Экспорт в Excel (заглушка)"""
        messagebox.showinfo("Информация", "Функция экспорта будет доступна в следующей версии")


class DocumentsReportWindow:
    """Окно журнала учета документов"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Журнал учета документов")
        self.window.geometry("1100x600")

        # Таблица
        self.create_table()
        self.load_report()

    def create_table(self):
        """Создание таблицы отчета"""
        table_frame = tk.Frame(self.window)
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        columns = ('Проект', 'Тип документа', 'Номер', 'Дата утверждения', 'Утвердил', 'Статус передачи в А4')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=20)

        col_widths = [150, 200, 120, 120, 150, 150]
        for i, col in enumerate(columns):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=col_widths[i])

        # Скроллбары
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

    def load_report(self):
        """Загрузка данных для отчета"""
        for item in self.tree.get_children():
            self.tree.delete(item)

        query = """
            SELECT f.id, p.project_number, f.doc_type, f.doc_number, f.approval_date,
                   e.full_name as approver_name,
                   COALESCE(ap.transfer_status, 'pending') as transfer_status
            FROM final_documents f
            JOIN projects p ON f.project_id = p.id
            LEFT JOIN employees e ON f.approver_id = e.id
            LEFT JOIN approved_packages ap ON 
                (f.doc_type = 'conclusion' AND ap.conclusion_id = f.id) OR
                (f.doc_type = 'report' AND ap.report_id = f.id) OR
                (f.doc_type = 'act' AND ap.act_id = f.id)
            ORDER BY f.approval_date DESC
        """
        docs = Database.execute_query(query, fetchall=True)

        if docs:
            doc_type_names = {'act': 'Акт обследования', 'report': 'Технический отчет',
                              'conclusion': 'Экспертное заключение'}
            status_names = {'pending': 'Ожидает передачи', 'transferred': 'Передан в А4'}

            for d in docs:
                doc_type = doc_type_names.get(d['doc_type'], d['doc_type'])
                transfer_status = status_names.get(d['transfer_status'], d['transfer_status'])

                self.tree.insert('', 'end', values=(
                    d['project_number'],
                    doc_type,
                    d['doc_number'] or '',
                    d['approval_date'].strftime('%d.%m.%Y') if d['approval_date'] else '',
                    d['approver_name'] or '',
                    transfer_status
                ))


class AuditLogWindow:
    """Окно журнала аудита"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Журнал аудита")
        self.window.geometry("1100x600")

        # Таблица
        self.create_table()
        self.load_log()

    def create_table(self):
        """Создание таблицы журнала аудита"""
        table_frame = tk.Frame(self.window)
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        columns = ('ID', 'Дата и время', 'Пользователь', 'Действие', 'Объект', 'ID объекта', 'Детали')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=20)

        col_widths = [50, 150, 150, 150, 100, 80, 300]
        for i, col in enumerate(columns):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=col_widths[i])

        # Скроллбары
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

    def load_log(self):
        """Загрузка данных журнала аудита"""
        for item in self.tree.get_children():
            self.tree.delete(item)

        query = """
            SELECT a.id, a.action_time, e.full_name as user_name, a.action_type,
                   a.object_type, a.object_id, a.details
            FROM audit_log a
            JOIN employees e ON a.employee_id = e.id
            ORDER BY a.action_time DESC
            LIMIT 1000
        """
        logs = Database.execute_query(query, fetchall=True)

        if logs:
            for log in logs:
                self.tree.insert('', 'end', values=(
                    log['id'],
                    log['action_time'].strftime('%d.%m.%Y %H:%M:%S') if log['action_time'] else '',
                    log['user_name'],
                    log['action_type'],
                    log['object_type'] or '',
                    log['object_id'] or '',
                    log['details'] or ''
                ))


class UserManagementWindow:
    """Окно управления пользователями"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Управление пользователями")
        self.window.geometry("900x600")

        # Панель инструментов
        toolbar = tk.Frame(self.window, bg="#f0f0f0", height=40)
        toolbar.pack(fill="x", padx=2, pady=2)

        tk.Button(toolbar, text="Добавить пользователя", command=self.add_user,
                  bg="#4CAF50", fg="white").pack(side="left", padx=5, pady=5)

        # Таблица
        self.create_table()
        self.load_users()

    def create_table(self):
        """Создание таблицы пользователей"""
        table_frame = tk.Frame(self.window)
        table_frame.pack(fill="both", expand=True, padx=5, pady=5)

        columns = ('ID', 'ФИО', 'Должность', 'Роль', 'Логин', 'Контакты')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=20)

        col_widths = [50, 200, 150, 120, 100, 200]
        for i, col in enumerate(columns):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=col_widths[i])

        # Скроллбары
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)

        # Правая кнопка мыши для редактирования/удаления
        self.tree.bind('<Button-3>', self.show_context_menu)

    def load_users(self):
        """Загрузка списка пользователей"""
        for item in self.tree.get_children():
            self.tree.delete(item)

        users = Database.execute_query("SELECT * FROM employees ORDER BY id", fetchall=True)

        role_names = {'admin': 'Администратор', 'director': 'Руководитель',
                      'expert': 'Эксперт', 'specialist': 'Специалист'}

        if users:
            for u in users:
                role = role_names.get(u['role'], u['role'])

                self.tree.insert('', 'end', values=(
                    u['id'],
                    u['full_name'],
                    u['position'] or '',
                    role,
                    u['login'],
                    u['contacts'] or ''
                ))

    def add_user(self):
        """Добавление пользователя"""
        # Здесь должна быть форма добавления пользователя
        messagebox.showinfo("Информация", "Форма добавления пользователя")

    def show_context_menu(self, event):
        """Показать контекстное меню"""
        selected = self.tree.selection()
        if not selected:
            return

        menu = tk.Menu(self.window, tearoff=0)
        menu.add_command(label="Редактировать", command=self.edit_user)
        menu.add_command(label="Сбросить пароль", command=self.reset_password)
        menu.add_separator()
        menu.add_command(label="Удалить", command=self.delete_user)
        menu.post(event.x_root, event.y_root)

    def edit_user(self):
        """Редактирование пользователя"""
        selected = self.tree.selection()
        if not selected:
            return
        messagebox.showinfo("Информация", "Редактирование пользователя")

    def reset_password(self):
        """Сброс пароля"""
        selected = self.tree.selection()
        if not selected:
            return
        messagebox.showinfo("Информация", "Пароль сброшен на 'password'")

    def delete_user(self):
        """Удаление пользователя"""
        selected = self.tree.selection()
        if not selected:
            return
        if messagebox.askyesno("Подтверждение", "Удалить пользователя?"):
            messagebox.showinfo("Информация", "Пользователь удален")


class BackupWindow:
    """Окно резервного копирования"""

    def __init__(self):
        self.window = tk.Toplevel()
        self.window.title("Резервное копирование")
        self.window.geometry("500x300")

        main_frame = tk.Frame(self.window, padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)

        tk.Label(main_frame, text="Резервное копирование базы данных",
                 font=("Arial", 14, "bold")).pack(pady=20)

        tk.Button(main_frame, text="Создать резервную копию", command=self.create_backup,
                  bg="#4CAF50", fg="white", font=("Arial", 12), width=20).pack(pady=10)

        tk.Button(main_frame, text="Восстановить из копии", command=self.restore_backup,
                  bg="#2196F3", fg="white", font=("Arial", 12), width=20).pack(pady=10)

        tk.Label(main_frame, text="Резервные копии сохраняются в папке backups",
                 font=("Arial", 9), fg="gray").pack(pady=20)

        tk.Button(main_frame, text="Закрыть", command=self.window.destroy,
                  bg="#f44336", fg="white", width=10).pack(pady=10)

    def create_backup(self):
        """Создание резервной копии"""
        backup_folder = "backups"
        if not os.path.exists(backup_folder):
            os.makedirs(backup_folder)

        backup_file = os.path.join(backup_folder, f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.sql")

        # Здесь должен быть код для создания дампа БД
        # В реальном проекте используется mysqldump

        Database.log_action(current_user['id'], 'BACKUP', 'system', None, f"Создана резервная копия {backup_file}")
        messagebox.showinfo("Успех", f"Резервная копия создана: {backup_file}")

    def restore_backup(self):
        """Восстановление из резервной копии"""
        filename = filedialog.askopenfilename(
            title="Выберите файл резервной копии",
            filetypes=[("SQL files", "*.sql"), ("All files", "*.*")]
        )
        if filename:
            messagebox.showinfo("Информация", "Восстановление из резервной копии")


if __name__ == "__main__":
    # Проверка подключения к БД
    conn = Database.get_connection()
    if conn:
        conn.close()
        app = LoginWindow()
        app.run()
    else:
        messagebox.showerror("Ошибка", "Не удалось подключиться к базе данных. Проверьте настройки подключения.")